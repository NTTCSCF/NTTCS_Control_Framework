import mimetypes
from time import sleep, time
from typing import Dict, Any

from django.contrib.auth import authenticate, login
from django.contrib.auth.decorators import login_required
from django.contrib.auth.mixins import LoginRequiredMixin
from django.contrib.auth.models import Group
from django.contrib.auth.views import LoginView

import json
import pandas as pd

from acounts.models import User
from django.core.paginator import Paginator
from django.shortcuts import render, HttpResponse, redirect
from datetime import datetime

from .models import Assessment, MaturirtyTable, AsociacionMarcos, Assessmentguardados, \
    NttcsCf20231, Domains, Evidencerequestcatalog, Evidencias, MapeoMarcos, AssessmentCreados, \
    AsociacionEvidenciasGenericas, AsociacionEvidenciasCreadas, TiposIniciativas, Iniciativas, \
    AssessmentEs, MaturirtyTableEs, EvidencerequestcatalogEs, Cliente, Proyecto, AsociacionUsuariosProyecto, \
    AsociacionProyectoAssessment, ProyectosMejora, AsociacionProyectoMejoraIniciativa, Entrevistas, \
    AsociacionEntrevistasUsuarios, AsociacionPlanProyectosProyectos, PlanProyectoMejora

from django.views.generic import TemplateView, ListView
import mysql.connector
from django.contrib import messages
import csv
from bs4 import BeautifulSoup

from docx import Document
from docx.shared import Cm


# Create your views here.

# Clase para la pagina del login
class index(LoginView):
    template_name = "homepage/index.html"

    def post(self, request):
        user = request.POST.get('user')
        pas = request.POST.get('pass')
        usuario = authenticate(request, username=user, password=pas)

        if usuario is not None:
            if usuario.last_login is None:
                login(request, usuario)
                return redirect('creacionPass')
            else:
                login(request, usuario)
                return redirect('menu')

        else:
            return render(request, self.template_name)


@login_required
def logout(request):
    logout(request)
    return redirect('/')


# funcion utilizada para comprobar que la contraseña intruducida es correcta
def contrasenaValida(password: str) -> bool:
    largo = False  # se pone a true si la cadena es lo suficientemente larga
    mayus = False  # se pone a true si la cadena contiene una mayuscula
    numerico = False  # se pone a true si la cadena contiene un numero
    caracterEspecial = False  # se pone a true si la cadena contiene un caracter especial /\.,:;!@#$%^&*()-_=+
    if len(password) > 8:  # comprobacion de longitud
        largo = True

    for i in password:
        if i.isupper():  # comprobacion de mayus
            mayus = True
        if i.isnumeric():  # comprobacion de numero
            numerico = True
        if i in '/\.,:;!@#$%^&*()-_=+':  # comprobacion de caracter especial
            caracterEspecial = True

    return largo and mayus and numerico and caracterEspecial  # retornemos and logico entre los 4 requerimientos


# Clase para la pagina de creacion de pass
class CreacionPass(LoginRequiredMixin, TemplateView):
    login_url = ""
    redirect_field_name = "redirect_to"

    template_name = "homepage/CreacionPass.html"

    def post(self, request, **knwargs):
        password = request.POST.get('password')  # valor del password
        passwordModificar = request.POST.get('passwordModificar')  # valor del passwordModificar
        password2Modificar = request.POST.get('password2Modificar')  # valor del password2Modificar
        user = request.user

        if password != '' and passwordModificar != '' and password2Modificar != '':  # comprobacion de entrega de cademas vacias
            if passwordModificar == password2Modificar:  # comprobacion de que las dos nuevas pass sean iguales
                if password != passwordModificar:  # coprobacion de que la contraseña nueva es distinta que la anterior
                    if user.check_password(password):  # comprobamos que la pass sea la correcta para el usuario
                        if contrasenaValida(
                                passwordModificar):  # comprobamos si la nueva pass cumple los requisitos de seguridad
                            user.set_password(passwordModificar)
                            user.save()
                            messages.success(request, 'La contraseña ha sido cambiada correctamente')
                            return redirect('logout')  # hacemos log out
                        else:
                            messages.error(request,
                                           'La contraseña debe contener, 8 caracteres, alguna mayuscula, algun caracter '
                                           'especial y algun numero')  # Se crea mensage de error
                    else:
                        messages.error(request, 'La contraseña no es correcta')  # Se crea mensage de error
                else:
                    messages.error(request,
                                   'La contraseña tiene que ser distinta a la anterior')  # Se crea mensage de error
            else:
                messages.error(request, 'Las contraseñas no coinciden')  # Se crea mensage de error
        else:
            messages.error(request, 'Debe rellenar todos los datos.')  # Se crea mensage de error
        return render(request, self.template_name)


# Clase para la pagina de Perfil
class Perfil(LoginRequiredMixin, TemplateView):
    login_url = ""
    redirect_field_name = "redirect_to"

    template_name = "homepage/perfil.html"

    def post(self, request, **knwargs):
        password = request.POST.get('password')  # valor del password
        passwordModificar = request.POST.get('passwordModificar')  # valor del passwordModificar
        password2Modificar = request.POST.get('password2Modificar')  # valor del password2Modificar
        user = request.user
        if password != '' and passwordModificar != '' and password2Modificar != '':  # comprobacion de entrega de cademas vacias
            if passwordModificar == password2Modificar:  # comprobacion de que las dos nuevas pass sean iguales
                if password != passwordModificar:  # coprobacion de que la contraseña nueva es distinta que la anterior
                    if user.check_password(password):  # comprobamos que la pass sea la correcta para el usuario
                        if contrasenaValida(
                                passwordModificar):  # comprobamos si la nueva pass cumple los requisitos de seguridad
                            user.set_password(passwordModificar)
                            user.save()
                            messages.success(request,
                                             'La contraseña ha sido cambiada correctamente')  # Se crea mensage de error
                        else:
                            messages.error(request,
                                           'La contraseña debe contener, 8 caracteres, alguna mayuscula, algun caracter '
                                           'especial y algun numero')  # Se crea mensage de error
                    else:
                        messages.error(request, 'La contraseña no es correcta')  # Se crea mensage de error
                else:
                    messages.error(request,
                                   'La contraseña tiene que ser distinta a la anterior')  # Se crea mensage de error
            else:
                messages.error(request, 'Las contraseñas no coinciden')  # Se crea mensage de error
        else:
            messages.error(request, 'Debe rellenar todos los datos.')  # Se crea mensage de error
        return render(request, self.template_name)


# Clase para la pagina de Usuarios
class Usuarios(LoginRequiredMixin, TemplateView):
    login_url = ""
    redirect_field_name = "redirect_to"
    template_name = "homepage/Usuarios.html"

    def get_context_data(self, **knwargs):
        context = super(Usuarios, self).get_context_data(**knwargs)
        context["grupos"] = Group.objects.all()  # contexto para el selector de gupos
        context["usuarios"] = User.objects.all()  # contexto para el selector de usuarios
        context["seleccionado"] = False  # variable para el despliegue de la modificacion de usuario
        return context

    def post(self, request, **knwargs):

        if 'boton3' in request.POST:  # if que recoge la pulsacion del boton de añadir usuario
            grupo = request.POST.get('selector')  # recogemos valor de selector de grupos
            nUsuario = request.POST.get('nUsuario')  # recogemos valor del nombre de usuario
            password = request.POST.get('password')  # recogemos el valor de la contraseña

            if grupo != 'None' and nUsuario != '' and password != '':  # comprobamos si todas las casillas esta rellenas.
                usuario = User.objects.create_user(username=nUsuario, password=password,
                                                   rol=grupo)  # creamos el usuario
                usuario.groups.add(Group.objects.get(name__in=[grupo]))  # le añadimos los permisos
                usuario.save()  # guardamos el usuario
            else:
                messages.error(request, 'Debe rellenar todos los datos.')  # Se crea mensage de error

        elif 'selector1' in request.POST:  # if para la funcionalidad del selector
            if request.POST.get('selector1') == 'None':  # comprobamos si se selecciona el No Seleccionado
                context = super(Usuarios, self).get_context_data(**knwargs)
                context["grupos"] = Group.objects.all()
                context["usuarios"] = User.objects.all()
                context["seleccionado"] = False  # fijamos el desplegable a false
                return render(request, self.template_name, context=context)
            else:
                context = super(Usuarios, self).get_context_data(**knwargs)
                context["grupos"] = Group.objects.all()
                context["usuarios"] = User.objects.all()
                context["seleccionado"] = True  # fijamos el desplegable a true para que se muestre la info del usuario.
                usuarioSeleccionado = User.objects.get(username=request.POST.get('selector1'))  # fijamos el usuario
                context["seleccion"] = usuarioSeleccionado  # pasamos el usuario
                context["grupoSeleccionado"] = usuarioSeleccionado.groups.all()[0]  # pasamos el grupo del usuario
                # seleccionado
                request.session["usuarioSeleccionado"] = usuarioSeleccionado.username  # guardamos en la sesion el
                # usuario que se ha seleccioando.
                print(usuarioSeleccionado.password)
                return render(request, self.template_name, context=context)
        elif 'selector3' in request.POST:  # Recogemos el valor de los botones de modificar y eliminar.
            boton1 = request.POST.get('boton1')  # valor del boton 1
            grupo = request.POST.get('selector3')  # valor del boton 1
            nUsuarioModificar = request.POST.get('nUsuarioModificar')  # valor del nUsuarioModificar
            passwordModificar = request.POST.get('passwordModificar')  # valor del passwordModificar
            user = User.objects.get(username=request.session.get('usuarioSeleccionado'))  # recogemos el usuario
            # seleccionado en la sesion.
            if grupo != 'None' and nUsuarioModificar != '' and nUsuarioModificar != '':  # comprobamos que se hayan
                # rellenado todos los campos
                if boton1 == 'btn1':  # recogemos la pulsacion del boton modificar
                    user.username = nUsuarioModificar  # fijamos el nombre del usuario
                    user.groups.clear()  # limpiamos los permisos anteriores
                    user.groups.add(Group.objects.get(name__in=[grupo]))  # añadimos los nuevos permisos, pueden ser
                    # los mismos
                    if passwordModificar.startswith('pbkdf2_sha256$'):  # comprobamos si el valor de la contraseña es
                        # la hasheada, osea la anterior, no podemos saber el valor sin hashear
                        user.password = passwordModificar
                    else:  # si no es el valor hasheado lo tenemos que encriptar
                        user.set_password(passwordModificar)  # para eso usamos el set_password
                    user.save()
                else:  # recogemos la pulsacion del boton Eliminar
                    user.delete()  # eliminamos el usuario seleccioando
            else:
                messages.error(request, 'Debe rellenar todos los datos.')  # Se crea mensage de error

        context = super(Usuarios, self).get_context_data(**knwargs)
        context["grupos"] = Group.objects.all()
        context["usuarios"] = User.objects.all()
        return render(request, self.template_name, context=context)


class assessment(LoginRequiredMixin, TemplateView):
    ''' Definición de la clase 'Assessment' '''

    login_url = ""
    redirect_field_name = "redirect_to"
    template_name = "homepage/assessment.html"
    # Se hace una conexión a la base de datos.
    conn = mysql.connector.connect(user='root', password="NTTCSCF2023", host='127.0.0.1', database='nttcs_cf',
                                   auth_plugin='mysql_native_password')

    def contextTotal(self, request, select, assSelect, context):

        assGuardado = Assessmentguardados.objects.get(id_assessment=assSelect)

        context["NombreAss"] = assSelect
        context["assessment"] = assGuardado
        context["assess"] = AssessmentCreados.objects.filter(assessment=assGuardado)
        control = AssessmentCreados.objects.get(assessment=assGuardado, control_id=select)

        if assGuardado.idioma == 'en':
            context["valMad"] = MaturirtyTable.objects.all()  # consulta para el desplegable de la valoracion de madurez
            context["evidenciasGenerricas"] = Evidencerequestcatalog.objects.all()
            lista = []
            for i in AsociacionEvidenciasGenericas.objects.filter(assessment=control):
                lista += [i.evidencia_id]
            context["listaEvidencias"] = lista
        else:
            context[
                "valMad"] = MaturirtyTableEs.objects.all()  # consulta para el desplegable de la valoracion de madurez
            context["evidenciasGenerricas"] = EvidencerequestcatalogEs.objects.all()

            lista = []
            for i in AsociacionEvidenciasGenericas.objects.filter(assessment=control):
                lista += [i.evidencia_id_es]
            context["listaEvidencias"] = lista

        request.session["controlSelect"] = select
        context["control"] = control
        context["criteriovaloracioncontexto"] = ['CCMM L0 \n No realizado',
                                                 'CCMM L1 \n Realizado de manera informal',
                                                 'CCMM L2 \n Planificado y rastreado',
                                                 'CCMM L3 \n Bien definido',
                                                 'CCMM L4 \n Controlado Cuantitativamente',
                                                 'CCMM L5 \n Mejorando Continuamente']

        context["criteriovaloracion"] = control.criteriovaloracion.split('[-33--33-]')
        context["evidencias"] = AsociacionEvidenciasGenericas.objects.filter(assessment=control)
        context["evidencias2"] = AsociacionEvidenciasCreadas.objects.filter(id_assessment=control)

        context["tiposIniciativas"] = TiposIniciativas.objects.all()
        context["iniciativas"] = Iniciativas.objects.all()

        evidenciasRecomendadas = Assessment.objects.get(id=select).evidence_request_references

        if evidenciasRecomendadas == '' or evidenciasRecomendadas is None:
            context["recomendacion"] = False
        else:

            evidenciasRecomendadas = evidenciasRecomendadas.split('\n')
            r = ''
            for i in evidenciasRecomendadas:
                if assGuardado.idioma == 'en':
                    evidencia = Evidencerequestcatalog.objects.get(evidence_request_references=i)
                    if not AsociacionEvidenciasGenericas.objects.filter(evidencia=evidencia, assessment=control):
                        r += i + ', '
                else:
                    evidencia = EvidencerequestcatalogEs.objects.get(evidence_request_references=i)
                    if not AsociacionEvidenciasGenericas.objects.filter(evidencia_id_es=evidencia, assessment=control):
                        r += i + ', '
            if r != '':
                context["recomendacion"] = True
                context["eviRecomendada"] = r[:len(r) - 2]
            else:
                context["recomendacion"] = False

        return request, context

    def get_context_data(self, **knwargs):
        '''El objetivo de este método es proporcionar datos de contexto para una vista,
               que luego se pueden utilizar en una plantilla HTML para renderizar la página web.'''

        # Guardamos el id (nombre) del assesment en una variable.
        assSelect = self.request.session.get('assessmentGuardado')
        # Se consigue de la bd el assessment correspondiente al id conseguido previamente.
        assGuardado = Assessmentguardados.objects.get(id_assessment=assSelect)
        # Actualizamos la fecha de la última modificación del assessment en formato YYYY/MM/DD.
        assGuardado.fecha_ultima_modificacion = datetime.now().isoformat().split('T')[0]
        # Se guardan los cambios en la bd.
        assGuardado.save()

        # Esto inicializa un diccionario llamado context con algunos datos de contexto.
        context = super(assessment, self).get_context_data(**knwargs)
        # Se guarda en el contexto el nombre del assessment (o la id).
        context["NombreAss"] = assSelect
        context["assess"] = AssessmentCreados.objects.filter(assessment=assGuardado)

        # Actualizaremos el contexto en base del idioma.
        if assGuardado.idioma == 'en':
            # Consulta para el desplegable de la valoración de madurez.
            context["valMad"] = MaturirtyTable.objects.all()
            # Se guardan las evidencias.
            context["evidenciasGenerricas"] = Evidencerequestcatalog.objects.all()
        else:
            # Consulta para el desplegable de la valoración de madurez.
            context["valMad"] = MaturirtyTableEs.objects.all()
            # Se guardan las evidencias.
            context["evidenciasGenerricas"] = EvidencerequestcatalogEs.objects.all()

        # Se guardan los tipos de inicitaivas.
        context["tiposIniciativas"] = TiposIniciativas.objects.all()
        # Se guardan las inicitaivas existentes.
        context["iniciativas"] = Iniciativas.objects.all()

        self.request.session["controlSelect"] = 'noSel'

        # Se devuelve el contexto.
        return context

    def post(self, request, **knwargs):
        ''' Método usado para manejar solcitudes HTTP POST enviadas por el cliente, en este caso,
        a través de un formulario. '''

        assSelect = request.session.get('assessmentGuardado')
        ''' Guardamos el id (nombre) del assesment en una variable. '''

        select = request.POST.get('selector')
        '''Valor de el selector de control.'''

        guardarValoracionBttn = request.POST.get('boton2')
        '''Valor del 'boton 2', que corresponde al botón de 'guardar valoración'.'''

        añadirEvidenciaBttn = request.POST.get('boton4')
        '''Valor del 'boton 4', que corresponde al botón de 'añadir evidencia'.'''

        seleccionarEvidenciaBttn = request.POST.get('boton5')
        '''Valor del 'boton 5', que corresponde al botón de 'seleccionar evidencia'.'''

        crearEvidenciaBttn = request.POST.get('boton6')
        '''Valor del 'boton 6', que corresponde al botón de 'crear iniciativa'.'''

        asignarIniciativaBttn = request.POST.get('boton7')
        '''Valor del 'boton 7', que corresponde al botón de 'asignar iniciativa'.'''

        boton8 = request.POST.get('boton8')
        '''Valor del 'boton 8', que corresponde al botón de 'añadir' cuando se recomiendan evidencias.'''

        btnEliminarEvidencia = request.POST.get('btnEliminarEvidencia')
        '''Valor del 'btnEliminarEvidencia'.'''

        # Se determina si se ha seleccionado el selector de marcos.
        if 'selector' in request.POST:
            # Si no se ha seleccionado ningún marco.
            if select == 'noSel':
                # Se coge el assessment guardado con el id del assessment seleccionado.
                assGuardado = Assessmentguardados.objects.get(id_assessment=assSelect)

                # Esto inicializa un diccionario llamado context con algunos datos de contexto.
                context = super(assessment, self).get_context_data(**knwargs)
                context["NombreAss"] = assSelect
                context["assess"] = AssessmentCreados.objects.filter(assessment=assGuardado)

                # Actualizaremos el contexto en base del idioma.
                if assGuardado.idioma == 'en':
                    # Consulta para el desplegable de la valoración de madurez.
                    context["valMad"] = MaturirtyTable.objects.all()
                    # Se guardan las evidencias.
                    context["evidenciasGenerricas"] = Evidencerequestcatalog.objects.all()
                else:
                    # Consulta para el desplegable de la valoración de madurez.
                    context["valMad"] = MaturirtyTableEs.objects.all()
                    # Se guardan las evidencias.
                    context["evidenciasGenerricas"] = EvidencerequestcatalogEs.objects.all()

                # Se guarda en una sessón del servidor el marco seleccionado.
                request.session["controlSelect"] = select

                # Se guardan los tipos de inicitaivas.
                context["tiposIniciativas"] = TiposIniciativas.objects.all()
                # Se guardan las inicitaivas existentes.
                context["iniciativas"] = Iniciativas.objects.all()

                # Se renderiza el template a base del contexto.
                return render(request, self.template_name, context=context)

            else:
                # Esto inicializa un diccionario llamado context con algunos datos de contexto.
                context = super(assessment, self).get_context_data(**knwargs)
                request, context = self.contextTotal(request, select, assSelect, context)

                # Se renderiza el template a base del contexto.
                return render(request, self.template_name, context=context)

        # Comprobamos si se ha pulsado el botón de 'guardar valoración'.
        elif guardarValoracionBttn == 'btn2':

            # Si se ha seleccionado ningún control.
            if request.session["controlSelect"] != 'noSel':
                # Conseguimos el assessment de la bd que se ha seleccionado anteriormente.
                assGuardado = Assessmentguardados.objects.get(id_assessment=assSelect)

                ''' Se consigue el assessment de la base de datos y se actualiza con la información
                corresponiente del formulario. '''
                control = AssessmentCreados.objects.get(assessment=assGuardado,
                                                        control_id=request.session["controlSelect"])
                control.respuesta = str(request.POST.get('respuesta'))
                control.valoracion = request.POST.get('valmad')
                control.valoracionobjetivo = request.POST.get('valmadob')
                control.save()

            # Si no se ha seleccionado ningún control.
            else:
                # Se crea un mensaje de error.
                messages.error(request,
                               'CONTROL INCORRECTO: Necesita seleccionar un control para realizar esta acción')

            # Esto inicializa un diccionario llamado context con algunos datos de contexto.
            context = super(assessment, self).get_context_data(**knwargs)
            request, context = self.contextTotal(request, request.session.get('controlSelect'), assSelect, context)

            # Se renderiza el template a base del contexto.
            return render(request, self.template_name, context=context)

        # Comprobamos si se ha pulsado el botón de 'guardar valoración'.
        elif añadirEvidenciaBttn == 'btn4':

            '''Tarjeta de Nueva Evidencia  '''
            # Se coge el input 'id referencia' de la nueva evidencia y se le concatena '-C', pues así
            # está guardada en la bd.
            idEvidencia = request.POST.get('idEvidencia') + '-C'
            # Valor del DescripcionEvidencia.
            descripcionEvidencia = request.POST.get('DescripcionEvidencia')
            # Valor del linkEvidencia.
            linkEvidencia = request.POST.get('linkEvidencia')

            # Se guarda el control seleccionado anteriormente.
            controlId = request.session["controlSelect"]

            # Si se ha seleccionado algún control.
            if controlId != 'noSel':
                # Si ni la id como la descripción de la evidencia están vacías.
                if idEvidencia != '' and descripcionEvidencia != '':
                    # Si la evidencia no existe en la base de datos, se puede crear.
                    if not Evidencias.objects.filter(evidencia_id=idEvidencia).exists():
                        # Se crea una nueva evidencia y se guarda en la BD.
                        ev = Evidencias(evidencia_id=idEvidencia, comentario=descripcionEvidencia, links=linkEvidencia,
                                        control_id=controlId,
                                        assessment=Assessmentguardados.objects.get(id_assessment=assSelect))
                        ev.save()

                        assGuardado = Assessmentguardados.objects.get(id_assessment=assSelect)

                        # Obtenemos el assesment creado a partir de la información obtenida.
                        control = AssessmentCreados.objects.get(assessment=assGuardado,
                                                                control_id=request.session["controlSelect"])
                        # Se guarda el input de 'respuesta'.
                        control.respuesta = str(request.POST.get('respuesta'))
                        # Se guarda el input de 'valoración de madurez'.
                        control.valoracion = request.POST.get('valmad')
                        # Se guarda el input de 'valoración de madurez objetivo'.
                        control.valoracionobjetivo = request.POST.get('valmadob')

                        # TODO: Missing control.save() ?
                        control.save()

                        # Se asocia la evidencia con el assessment creado.

                        evidencia = AsociacionEvidenciasCreadas(id_evidencia=ev, id_assessment=control)
                        # Se guarda en la BD.
                        evidencia.save()

                        # Esto inicializa un diccionario llamado context con algunos datos de contexto.
                        context = super(assessment, self).get_context_data(**knwargs)
                        request, context = self.contextTotal(request, controlId, assSelect, context)

                        # Renderiza el template en base el contexto.
                        return render(request, self.template_name, context=context)
                    # Si la evidencia  existe en la base de datos, nos se puede crear.
                    else:
                        # Esto inicializa un diccionario llamado context con algunos datos de contexto.
                        context = super(assessment, self).get_context_data(**knwargs)
                        request, context = self.contextTotal(request, controlId, assSelect, context)
                        # Se crea mensaje de error.
                        messages.error(request, 'EVIDENCIA INCORRECTA: La evidencia introducida ya existe')

                        # Renderiza el template en base el contexto.
                        return render(request, self.template_name, context=context)
                # Si la id o la descripción de la evidencia están vacías.
                else:
                    # Esto inicializa un diccionario llamado context con algunos datos de contexto.
                    context = super(assessment, self).get_context_data(**knwargs)
                    request, context = self.contextTotal(request, controlId, assSelect, context)
                    # Se crea mensaje de error.
                    messages.error(request, 'EVIDENCIA INCORRECTA Necesita introducir un id y una descripcion para la '
                                            'evidencia')

                    # Renderiza el template en base el contexto.
                    return render(request, self.template_name, context=context)
            # Si no se ha seleccionado ningñun control.
            else:
                # Se crea un mensaje de error.
                messages.error(request,
                               'CONTROL INCORRECTO: Necesita seleccionar un control para realizar esta acción')

                # Esto inicializa un diccionario llamado context con algunos datos de contexto.
                context = super(assessment, self).get_context_data(**knwargs)
                # Se guarda el assessment de la tabla 'Assessmentguardados' en base del id del assessment seleccionado.
                assGuardado = Assessmentguardados.objects.get(id_assessment=assSelect)

                # Guardamos en el contexto el identificador del assessment que, a su vez, es el nombre.
                context["NombreAss"] = assSelect
                # Se guarda en el contexto el assessment creado en base de la variable anterior.
                context["assess"] = AssessmentCreados.objects.filter(assessment=assGuardado)

                # Actualizaremos el contexto en base del idioma.
                if assGuardado.idioma == 'en':
                    # Consulta para el desplegable de la valoración de madurez.
                    context["valMad"] = MaturirtyTable.objects.all()
                    # Se guardan las evidencias.
                    context["evidenciasGenerricas"] = Evidencerequestcatalog.objects.all()
                else:
                    # Consulta para el desplegable de la valoración de madurez.
                    context["valMad"] = MaturirtyTableEs.objects.all()
                    # Se guardan las evidencias.
                    context["evidenciasGenerricas"] = EvidencerequestcatalogEs.objects.all()

                # Se guardan los tipos de inicitaivas.
                context["tiposIniciativas"] = TiposIniciativas.objects.all()
                # Se guardan las inicitaivas existentes
                context["iniciativas"] = Iniciativas.objects.all()

                # Se renderizará el template en base del contexto.
                return render(request, self.template_name, context=context)

        # Detecta si se ha presionado el botón de 'seleccionar evidencia'.
        elif seleccionarEvidenciaBttn == 'btn5':

            # Si se ha seleccionado algún control.
            if request.session["controlSelect"] != 'noSel':
                # Se reoge el valor del selector de evidencia generica.
                selectorEvidencia = request.POST.get('selectorEvidencia')

                # Si se ha seleccionado alguna evidencia.
                if selectorEvidencia != 'noSel':
                    # Se elige el assessment guardado en base del nombre del assessment.
                    assGuardado = Assessmentguardados.objects.get(id_assessment=assSelect)
                    #
                    control = AssessmentCreados.objects.get(assessment=assGuardado,
                                                            control_id=request.session["controlSelect"])
                    # Si el idioma del assessment esta en ingles.
                    if assGuardado.idioma == 'en':
                        '''Se actualiza la tabla de 'AsociacionEvidenciasGenericas'. '''
                        evidencia = Evidencerequestcatalog.objects.get(evidence_request_references=selectorEvidencia)
                        eviGenerica = AsociacionEvidenciasGenericas(evidencia=evidencia, assessment=control)
                        eviGenerica.save()
                    else:
                        '''Se actualiza la tabla de 'AsociacionEvidenciasGenericas'. '''
                        evidencia = EvidencerequestcatalogEs.objects.get(evidence_request_references=selectorEvidencia)
                        eviGenerica = AsociacionEvidenciasGenericas(evidencia_id_es=evidencia, assessment=control)
                        eviGenerica.save()

                    # Se elige el assessment guardado en base del nombre del assessment.
                    assGuardado = Assessmentguardados.objects.get(id_assessment=assSelect)
                    control = AssessmentCreados.objects.get(assessment=assGuardado,
                                                            control_id=request.session["controlSelect"])
                    # Se guarda el input de 'respuesta'.
                    control.respuesta = str(request.POST.get('respuesta'))
                    # Se guarda el input de 'valoración de madurez'.
                    control.valoracion = request.POST.get('valmad')
                    # Se guarda el input de 'valoración de madurez objetivo'.
                    control.valoracionobjetivo = request.POST.get('valmadob')
                    # Se guarda en la bd.
                    control.save()
                    # Esto inicializa un diccionario llamado context con algunos datos de contexto.
                    context = super(assessment, self).get_context_data(**knwargs)
                    request, context = self.contextTotal(request, request.session["controlSelect"], assSelect, context)

                    # Se renderiza la página en base del contexto.
                    return render(request, self.template_name, context=context)

                # Si no se ha seleccionado ninguna evidencia.
                else:
                    # Esto inicializa un diccionario llamado context con algunos datos de contexto.
                    context = super(assessment, self).get_context_data(**knwargs)
                    request, context = self.contextTotal(request, request.session["controlSelect"], assSelect, context)

                    # Se crea un mensaje de error.
                    messages.error(request, 'EVIDENCIA INCORRECTA: Necesita seleccionar una evidencia')

                    # Se renderiza la pagina en base del contexto.
                    return render(request, self.template_name, context=context)

            # Si no se ha sleccionado ningún control.
            else:
                # Se crea un mensaje de error
                messages.error(request, 'CONTROL INCORRECTO: Necesita seleccionar un control para realizar esta acción')
                # Esto inicializa un diccionario llamado context con algunos datos de contexto.
                context = super(assessment, self).get_context_data(**knwargs)
                # Se elige el assessment guardado en base del nombre del assessment.
                assGuardado = Assessmentguardados.objects.get(id_assessment=assSelect)

                # Guardamos en el contexto el identificador del assessment que, a su vez, es el nombre.
                context["NombreAss"] = assSelect
                # Se guarda en el contexto el assessment creado en base de la variable anterior.
                context["assess"] = AssessmentCreados.objects.filter(assessment=assGuardado)

                # Actualizamos el contexto en función del idioma del assessment.
                if assGuardado.idioma == 'en':
                    # Consulta para el desplegable de la valoración de madurez.
                    context["valMad"] = MaturirtyTable.objects.all()
                    # Se guardan las evidencias.
                    context["evidenciasGenerricas"] = Evidencerequestcatalog.objects.all()
                else:
                    # Consulta para el desplegable de la valoración de madurez.
                    context["valMad"] = MaturirtyTableEs.objects.all()
                    # Se guardan las evidencias.
                    context["evidenciasGenerricas"] = EvidencerequestcatalogEs.objects.all()

                # Se guardan los tipos de inicitaivas.
                context["tiposIniciativas"] = TiposIniciativas.objects.all()
                # Se guardan las inicitaivas existentes.
                context["iniciativas"] = Iniciativas.objects.all()

                # Se renderiza el template en función del contexto.
                return render(request, self.template_name, context=context)

        # Se comrpueba si se le ha dado al botón 'crear iniciativa'.
        elif crearEvidenciaBttn == 'btn6':

            selectorEvidencia = request.POST.get('selectorEvidenciaIniciativa')
            ''' Se recoge el valor del desplegable 'seleccione la evidencia'. '''

            nombreIniciativa = request.POST.get('nombreIniciativa')
            ''' Se recoge el valor del input 'nombre de la iniciativa'. '''

            DescripcionIniciativa = request.POST.get('DescripcionIniciativa')
            ''' Se recoge el valor del input 'descripción de la inciativa'. '''

            SelectorIniciativa = request.POST.get('SelectorIniciativa')
            ''' Se recoge el valor del desplegable 'seleccione el tipo de iniciativa. '''

            # Si se ha seleccionado un control.
            if request.session["controlSelect"] != 'noSel':
                # Si se ha seleccionado una evidencia.
                if selectorEvidencia != 'noSel':
                    # Si la iniciativa no existe, se crea en la bd.
                    if not Iniciativas.objects.filter(nombre=nombreIniciativa).exists():
                        # Si la evidencia introducida existe en la tabla 'Evidencerequestcatalog'.
                        if Evidencerequestcatalog.objects.filter(
                                evidence_request_references=selectorEvidencia).exists():

                            # Conseguimos el assessment de la tabla 'AssessmentCreados' en base del id del nombre del assessment.
                            assGuardado = Assessmentguardados.objects.get(id_assessment=assSelect)
                            control = AssessmentCreados.objects.get(assessment=assGuardado,
                                                                    control_id=request.session["controlSelect"])

                            '''Dependiendo del idioma del assessment guardamos la evidencia en un idioma u otro.'''
                            if assGuardado.idioma == 'en':
                                evidencia = Evidencerequestcatalog.objects.get(
                                    evidence_request_references=selectorEvidencia)
                                asociacion = AsociacionEvidenciasGenericas.objects.get(evidencia=evidencia,
                                                                                       assessment=control)
                            else:
                                evidencia = EvidencerequestcatalogEs.objects.get(
                                    evidence_request_references=selectorEvidencia)
                                asociacion = AsociacionEvidenciasGenericas.objects.get(evidencia_id_es=evidencia,
                                                                                       assessment=control)

                            # Se coge el tipo de la iniciativa en base del introducido en el desplegable.
                            tipoIniciativa = TiposIniciativas.objects.get(tipo=SelectorIniciativa)
                            # Añadimos la relación entre la tabla 'iniciativa' y 'tipos_iniciativas'.
                            iniciativa = Iniciativas(nombre=nombreIniciativa, descripcion=DescripcionIniciativa,
                                                     tipo=tipoIniciativa)
                            # Guardamos la inciativa en la BD.
                            iniciativa.save()

                            # Asociamos la tabla 'AsociacionEvidenciasGenericas' con la tabla 'iniciativa'.
                            asociacion.iniciativa = iniciativa
                            asociacion.save()

                            # Actualizamos la tabla 'assessment_creados' por los datos introducidos.
                            control.respuesta = str(request.POST.get('respuesta'))
                            control.valoracion = request.POST.get('valmad')
                            control.valoracionobjetivo = request.POST.get('valmadob')
                            control.save()

                            # Esto inicializa un diccionario llamado context con algunos datos de contexto.
                            context = super(assessment, self).get_context_data(**knwargs)
                            request, context = self.contextTotal(request, request.session["controlSelect"], assSelect,
                                                                 context)

                            # Se renderiza el template en base del contexto.
                            return render(request, self.template_name, context=context)
                        # Si la evidencia no existe en la tabla 'Evidencerequestcatalog'.
                        else:
                            # Se consigue le evidencia de la base de datos.
                            evidencia = Evidencias.objects.get(evidencia_id=selectorEvidencia)

                            # Se relaciona el elemento de la tabla 'Assessmentguardados' con el de la tabla 'AssessmentCreados'
                            assGuardado = Assessmentguardados.objects.get(id_assessment=assSelect)
                            control = AssessmentCreados.objects.get(assessment=assGuardado,
                                                                    control_id=request.session["controlSelect"])
                            # Se asocia la evidencia con el control en la bd.
                            asociacion = AsociacionEvidenciasCreadas.objects.get(id_evidencia=evidencia,
                                                                                 id_assessment=control)
                            # Se relaciona la tabla 'tipoIniciativa' con 'Iniciativas' y se guarda en la BD.
                            tipoIniciativa = TiposIniciativas.objects.get(tipo=SelectorIniciativa)
                            iniciativa = Iniciativas(nombre=nombreIniciativa, descripcion=DescripcionIniciativa,
                                                     tipo=tipoIniciativa)
                            iniciativa.save()

                            # Se asocia el assessment con la iniciativa y se guarda en la BD.
                            asociacion.iniciativa = iniciativa
                            asociacion.save()

                            # Esto inicializa un diccionario llamado context con algunos datos de contexto.
                            context = super(assessment, self).get_context_data(**knwargs)
                            request, context = self.contextTotal(request, request.session["controlSelect"], assSelect,
                                                                 context)

                            # Se renderiza el template en base del contexto.
                            return render(request, self.template_name, context=context)

                    # Si la iniciativa existe, salta un error.
                    else:
                        # Esto inicializa un diccionario llamado context con algunos datos de contexto.
                        context = super(assessment, self).get_context_data(**knwargs)
                        request, context = self.contextTotal(request, request.session["controlSelect"], assSelect,
                                                             context)
                        # Mensaje de error.
                        messages.error(request,
                                       'INICIATIVA INCORRECTA: El nombre introducido ya esta en uso')

                        # Se renderiza el template en base del contexto.
                        return render(request, self.template_name, context=context)
                # Si no se ha seleccionado una evidencia.
                else:
                    # Esto inicializa un diccionario llamado context con algunos datos de contexto.
                    context = super(assessment, self).get_context_data(**knwargs)
                    request, context = self.contextTotal(request, request.session["controlSelect"], assSelect, context)
                    # Mensaje de error.
                    messages.error(request,
                                   'EVIDENCIA INCORRECTA: Necesita seleccionar una evidencia')

                    # Se renderiza el template en base del contexto.
                    return render(request, self.template_name, context=context)

            # Si no se ha seleccionado ningún control.
            else:
                # Mensaje de error.
                messages.error(request,
                               'CONTROL INCORRECTO: Necesita seleccionar un control para realizar esta acción')

                # Esto inicializa un diccionario llamado context con algunos datos de contexto.
                context = super(assessment, self).get_context_data(**knwargs)
                assGuardado = Assessmentguardados.objects.get(id_assessment=assSelect)

                # Guardamos en el contexto el identificador del assessment que, a su vez, es el nombre.
                context["NombreAss"] = assSelect
                # Se guarda en el contexto el assessment creado en base de la variable anterior.
                context["assess"] = AssessmentCreados.objects.filter(assessment=assGuardado)

                # Actualizamos el contexto en función del idioma del assessment.
                if assGuardado.idioma == 'en':
                    # Consulta para el desplegable de la valoración de madurez.
                    context["valMad"] = MaturirtyTable.objects.all()
                    # Se guardan las evidencias.
                    context["evidenciasGenerricas"] = Evidencerequestcatalog.objects.all()
                else:
                    # Consulta para el desplegable de la valoración de madurez.
                    context["valMad"] = MaturirtyTableEs.objects.all()
                    # Se guardan las evidencias.
                    context["evidenciasGenerricas"] = EvidencerequestcatalogEs.objects.all()

                # Se guardan los tipos de inicitaivas.
                context["tiposIniciativas"] = TiposIniciativas.objects.all()
                # Se guardan las inicitaivas existentes.
                context["iniciativas"] = Iniciativas.objects.all()

                # Se renderiza el template en base del contexto.
                return render(request, self.template_name, context=context)
        # Si se da al botón de 'Asignar Iniciativa'.
        elif asignarIniciativaBttn == 'btn7':
            selectEviasig = request.POST.get('selectEviasig')
            ''' Selector de evidencia. '''

            selectIniAsig = request.POST.get('selectIniAsig')
            ''' Selector de iniciativa. '''

            if request.session["controlSelect"] != 'noSel':
                if selectEviasig != 'noSel':
                    if selectIniAsig != 'noSel':
                        if Evidencerequestcatalog.objects.filter(
                                evidence_request_references=selectEviasig).exists():

                            assGuardado = Assessmentguardados.objects.get(id_assessment=assSelect)
                            control = AssessmentCreados.objects.get(assessment=assGuardado,
                                                                    control_id=request.session["controlSelect"])

                            if assGuardado.idioma == 'en':
                                evidencia = Evidencerequestcatalog.objects.get(
                                    evidence_request_references=selectEviasig)
                                asociacion = AsociacionEvidenciasGenericas.objects.get(evidencia=evidencia,
                                                                                       assessment=control)
                            else:
                                evidencia = EvidencerequestcatalogEs.objects.get(
                                    evidence_request_references=selectEviasig)
                                asociacion = AsociacionEvidenciasGenericas.objects.get(evidencia_id_es=evidencia,
                                                                                       assessment=control)

                            iniciativa = Iniciativas.objects.get(nombre=selectIniAsig)
                            asociacion.iniciativa = iniciativa
                            asociacion.save()

                            control.respuesta = str(request.POST.get('respuesta'))
                            control.valoracion = request.POST.get('valmad')
                            control.valoracionobjetivo = request.POST.get('valmadob')
                            control.save()
                            context = super(assessment, self).get_context_data(**knwargs)
                            request, context = self.contextTotal(request, request.session["controlSelect"], assSelect,
                                                                 context)
                            return render(request, self.template_name, context=context)
                        else:
                            evidencia = Evidencias.objects.get(evidencia_id=selectEviasig)
                            assGuardado = Assessmentguardados.objects.get(id_assessment=assSelect)
                            control = AssessmentCreados.objects.get(assessment=assGuardado,
                                                                    control_id=request.session["controlSelect"])
                            asociacion = AsociacionEvidenciasCreadas.objects.get(id_evidencia=evidencia,
                                                                                 id_assessment=control)

                            iniciativa = Iniciativas.objects.get(nombre=selectIniAsig)
                            asociacion.iniciativa = iniciativa
                            asociacion.save()
                            context = super(assessment, self).get_context_data(**knwargs)
                            request, context = self.contextTotal(request, request.session["controlSelect"], assSelect,
                                                                 context)
                            return render(request, self.template_name, context=context)
                    else:
                        context = super(assessment, self).get_context_data(**knwargs)
                        request, context = self.contextTotal(request, request.session["controlSelect"], assSelect,
                                                             context)
                        messages.error(request,
                                       'INICIATIVA INCORRECTA: Necesita seleccionar una iniciativa')  # Se crea mensage de error
                        return render(request, self.template_name, context=context)
                else:
                    context = super(assessment, self).get_context_data(**knwargs)
                    request, context = self.contextTotal(request, request.session["controlSelect"], assSelect,
                                                         context)
                    messages.error(request,
                                   'EVIDENCIA INCORRECTA: Necesita seleccionar una evidencia')  # Se crea mensage de error
                    return render(request, self.template_name, context=context)
            else:
                messages.error(request,
                               'CONTROL INCORRECTO: Necesita seleccionar un control para realizar esta acción')  # Se crea mensage de error
                context = super(assessment, self).get_context_data(**knwargs)
                assGuardado = Assessmentguardados.objects.get(id_assessment=assSelect)
                context["NombreAss"] = assSelect
                context["assess"] = AssessmentCreados.objects.filter(assessment=assGuardado)

                if assGuardado.idioma == 'en':
                    context[
                        "valMad"] = MaturirtyTable.objects.all()  # consulta para el desplegable de la valoracion de madurez
                    context["evidenciasGenerricas"] = Evidencerequestcatalog.objects.all()
                else:
                    context[
                        "valMad"] = MaturirtyTableEs.objects.all()  # consulta para el desplegable de la valoracion de madurez
                    context["evidenciasGenerricas"] = EvidencerequestcatalogEs.objects.all()

                context["tiposIniciativas"] = TiposIniciativas.objects.all()
                context["iniciativas"] = Iniciativas.objects.all()
                return render(request, self.template_name, context=context)
        elif boton8 == 'btn8':  # if encargado de rellenar las evidencias
            evidenciasRecomendadas = Assessment.objects.get(
                id=request.session["controlSelect"]).evidence_request_references.split('\n')
            assGuardado = Assessmentguardados.objects.get(id_assessment=assSelect)

            for i in evidenciasRecomendadas:
                control = AssessmentCreados.objects.get(assessment=assGuardado,
                                                        control_id=request.session["controlSelect"])
                if assGuardado.idioma == 'en':
                    evidencia = Evidencerequestcatalog.objects.get(
                        evidence_request_references=i)
                    if not AsociacionEvidenciasGenericas.objects.filter(evidencia=evidencia, assessment=control):
                        eviGenerica = AsociacionEvidenciasGenericas(evidencia=evidencia, assessment=control)
                        eviGenerica.save()
                else:
                    evidencia = EvidencerequestcatalogEs.objects.get(
                        evidence_request_references=i)
                    if not AsociacionEvidenciasGenericas.objects.filter(evidencia_id_es=evidencia, assessment=control):
                        eviGenerica = AsociacionEvidenciasGenericas(evidencia_id_es=evidencia, assessment=control)
                        eviGenerica.save()

            control = AssessmentCreados.objects.get(assessment=assGuardado,
                                                    control_id=request.session["controlSelect"])
            control.respuesta = str(request.POST.get('respuesta'))
            control.valoracion = request.POST.get('valmad')
            control.valoracionobjetivo = request.POST.get('valmadob')
            control.save()
            context = super(assessment, self).get_context_data(**knwargs)
            request, context = self.contextTotal(request, request.session["controlSelect"], assSelect, context)
            return render(request, self.template_name, context=context)
        elif btnEliminarEvidencia != '':  # if encargado de Eliminar las evidencias
            evidencia = AsociacionEvidenciasGenericas.objects.get(id=btnEliminarEvidencia)
            evidencia.delete()
            context = super(assessment, self).get_context_data(**knwargs)
            request, context = self.contextTotal(request, request.session["controlSelect"], assSelect, context)
            return render(request, self.template_name, context=context)


class assessmentselect(LoginRequiredMixin, TemplateView):
    ''' Definición de la clase 'assessmentselect' '''

    login_url = ""
    redirect_field_name = "redirect_to"
    template_name = "homepage/assessmentselect.html"

    def get_context_data(self, **knwargs):
        '''El objetivo de este método es proporcionar datos de contexto para una vista,
        que luego se pueden utilizar en una plantilla HTML para renderizar la página web.'''

        '''Asegura que cualquier funcionalidad definida en las clases base 'LoginRequiredMixin' y 'TemplateView'
        se ejecute antes de personalizarla en la subclase assessmentselect.'''
        context = super(assessmentselect, self).get_context_data(**knwargs)

        context["proyectos"] = AsociacionUsuariosProyecto.objects.filter(usuario=self.request.user)
        context["marcos"] = AsociacionMarcos.objects.all()
        return context

    def post(self, request, **knwargs):
        '''Método usado para manejar solcitudes HTTP POST enviadas por el cliente, en este caso,
        a través de un formulario.'''

        # Valor del input de nombre.
        nombre = request.POST.get('in')
        # Valor de selector de proyecto.
        selectorProyecto = request.POST.get('selectorProyecto')
        # Valor de selector de marcos de seguridad.
        select2 = request.POST.getlist('selector2')
        # Valor del input de idioma.
        idioma = request.POST.get('idioma')
        # Valor asociado al botón de editar (el id del assesmment asociado).
        btnEditar = request.POST.get('btnEditar')
        # Valor asociado al botón de archivar (el id del assesmment asociado).
        btnArchivar = request.POST.get('btnArchivar')

        # Comprueba si el usuario ha seleccionado un proyecto en el formulario.
        if 'selectorProyecto' in request.POST:

            '''Inicialización del contexto '''
            # Esto inicializa un diccionario llamado context con algunos datos de contexto.
            context = super(assessmentselect, self).get_context_data(**knwargs)
            # Añade al conexto los proyectos asociados al usuario.
            context["proyectos"] = AsociacionUsuariosProyecto.objects.filter(usuario=self.request.user)
            context["proyectoSelec"] = selectorProyecto
            context["proyectoSeleccionado"] = True
            # Recupera todos los objetos de la classe AsociacionProyectoAssessment filtrados por el proyecto
            # seleccionado en el selector de proyectos.
            context["assess"] = AsociacionProyectoAssessment.objects.filter(
                proyecto=Proyecto.objects.get(codigo=selectorProyecto), assessment__archivado=0)
            # Se guardan todos los marcos de seguridad en el contexto.
            context["marcos"] = AsociacionMarcos.objects.all()

            # Se guarda el proyecto seleccionado en una sessión en el servidor,
            request.session["proyectoSeleccionado"] = selectorProyecto

            # Renderiza el template en base del contexto.
            return render(request, self.template_name, context=context)

        # Si se decide editar algún assessments.
        elif 'btnEditar' in request.POST:
            request.session["assessmentGuardado"] = btnEditar

            # Redirección hacia la url especificada en base del nombre.
            return redirect("assessment")

        # Si se presiona el botón 'plan de proyecto' asociado a cada assessment.
        elif 'btnPlan' in request.POST:
            request.session["assessmentGuardado"] = request.POST.get('btnPlan')

            # Redirección hacia la url especificada en base del nombre.
            return redirect("planProyecto")

        # Si se presiona el botón 'archivar' asociado a cada assessment.
        elif 'btnArchivar' in request.POST:

            # Buscamos el assessment en función del id.
            ass = Assessmentguardados.objects.get(id_assessment=btnArchivar)
            # Se archiva el assessment (0 = no archivado / 1 = archivado).
            ass.archivado = 1
            # Se actualiza la fecha en formato 'AAAA-MM-DD'.
            ass.fecha_cierre = datetime.now().isoformat().split('T')[0]
            # Se guardan todos los cambios.
            ass.save()

            '''Inicialización del contexto '''
            # Se inicializa un diccionario llamado context con algunos datos de contexto.
            context = super(assessmentselect, self).get_context_data(**knwargs)
            # Conseguimos los proyectos asociados al usuario.
            context["proyectos"] = AsociacionUsuariosProyecto.objects.filter(usuario=self.request.user)
            context["proyectoSeleccionado"] = True
            # Guardamos en el contexto el proyecto seleccionado, el cual fue guardado anteriormente.
            context["proyectoSelec"] = request.session.get('proyectoSeleccionado')
            # Obtenemos todos los filas de la tabla 'AsociacionProyectoAssessment' que cumplan dichas condiciones.
            context["assess"] = AsociacionProyectoAssessment.objects.filter(
                proyecto=Proyecto.objects.get(codigo=request.session.get('proyectoSeleccionado')),
                assessment__archivado=0)
            # Se guardan todos los marcos de seguridad en el contexto.
            context["marcos"] = AsociacionMarcos.objects.all()

            # Renderiza el template en base del contexto.
            return render(request, self.template_name, context=context)

        # Si se escribe algo en el campo de nombre del assessment.
        elif 'in' in request.POST:

            # Si tanto el nombre como el selector de marcos están seleccionados.
            if nombre != '' and select2 != None:
                # Comprobamos que el nombre del assessment no existe en la BD (nombre nuevo)
                if Assessmentguardados.objects.filter(id_assessment=nombre).exists() == False:

                    # Creamos una nueva fila en 'assessmentguardados' inicializada con el nombre del assessment,
                    # la fecha de creación en formato 'AAAA-MM-DD', el idioma y el estado (0 = no archivado).
                    assessmentNuevo = Assessmentguardados(id_assessment=nombre,
                                                          archivado=0,
                                                          fecha_creacion=datetime.now().isoformat().split('T')[0],
                                                          idioma=idioma)
                    # Guardamos dicho assessment en la base de datos.
                    assessmentNuevo.save()

                    marcos = ''
                    ''' Marcos de NTT separados por un espacio.'''

                    marc = []
                    '''Lista de marcos de NTT.'''

                    # Iteramos sobre el selector de marcos.
                    for i in select2:
                        # Conseguimos el nombre de la tabla asociada a la iteración.
                        consulta = AsociacionMarcos.objects.get(marco_id=i).nombre_tabla

                        # Accede a la columna ('consulta') y accede a la fila con valor a 1,
                        # que pertenece al marco de NTT correspondiente.
                        c = MapeoMarcos.objects.extra(
                            where=[consulta + "='1'"])

                        # Recorremos la tabla del marco.
                        for fila in c:
                            if fila.ntt_id not in marc:
                                # Guardamos los marcos de NTT sin repetirlos.
                                marc += [fila.ntt_id]

                    for marco in marc:
                        # Añadimos un elemento de la lista al string, separado por un espacio.
                        marcos += marco + '\n'

                        # Obtenemos el assesment en función del idioma y del marco de NTT.
                        if idioma == 'en':
                            consulta = Assessment.objects.get(id=marco)
                        else:
                            consulta = AssessmentEs.objects.get(id=marco)

                        # Manera compacta de guardar muchos campos distintos en una sola variable.
                        # Posteriormente, se usará "split('-33--33-')" para acceder a cada campo.
                        criterioVal = consulta.campo9 + '[-33--33-]' + consulta.campo10 + '[-33--33-]' + consulta.campo11 + '[-33--33-]' + consulta.campo12 + '[-33--33-]' + consulta.campo13 + '[-33--33-]' + consulta.campo14

                        # Se crea una fila en 'AssessmentCreados' y se guarda en la BD.
                        a = AssessmentCreados(assessment=assessmentNuevo, control_id=str(marco),
                                              control_name=consulta.control,
                                              descripcion=consulta.control_description,
                                              pregunta=consulta.control_question,
                                              criteriovaloracion=criterioVal)
                        a.save()

                    # Actualizamos la fila creada anteriormente con los marcos de NTT.
                    assessmentNuevo.marcos = marcos
                    assessmentNuevo.save()

                    # Asociamos el assessment nuevo con el proyecto seleccionado y se guarda en la BD.
                    proyecto = Proyecto.objects.get(codigo=request.session.get('proyectoSeleccionado'))
                    asociacion = AsociacionProyectoAssessment(assessment=assessmentNuevo, proyecto=proyecto)
                    asociacion.save()

                    '''Inicialización del contexto '''
                    # Se inicializa un diccionario llamado context con algunos datos de contexto.
                    context = super(assessmentselect, self).get_context_data(**knwargs)
                    # Conseguimos los proyectos asociados al usuario.
                    context["proyectos"] = AsociacionUsuariosProyecto.objects.filter(usuario=self.request.user)
                    context["proyectoSeleccionado"] = True
                    # Guardamos en el contexto el proyecto seleccionado, el cual fue guardado anteriormente.
                    context["proyectoSelec"] = request.session.get('proyectoSeleccionado')
                    # Obtenemos todos los filas de la tabla 'AsociacionProyectoAssessment' que cumplan dichas condiciones.
                    context["assess"] = AsociacionProyectoAssessment.objects.filter(
                        proyecto=Proyecto.objects.get(codigo=request.session.get('proyectoSeleccionado')),
                        assessment__archivado=0)
                    # Se guardan todos los marcos de seguridad en el contexto.
                    context["marcos"] = AsociacionMarcos.objects.all()

                    # Se renderiza el template en base del contexto.
                    return render(request, self.template_name, context=context)

                # Si el nombre del assessment ya existe en la BD (nombre repetido).
                else:
                    # Mensaje de error.
                    messages.error(request, 'El Assessment ya exsiste.')

                    '''Inicialización del contexto '''
                    # Se inicializa un diccionario llamado context con algunos datos de contexto.
                    context = super(assessmentselect, self).get_context_data(**knwargs)
                    # Conseguimos los proyectos asociados al usuario.
                    context["proyectos"] = AsociacionUsuariosProyecto.objects.filter(usuario=self.request.user)
                    # Guardamos en el contexto el proyecto seleccionado, el cual fue guardado anteriormente.
                    context["proyectoSelec"] = request.session.get('proyectoSeleccionado')
                    context["proyectoSeleccionado"] = True
                    # Obtenemos todos los filas de la tabla 'AsociacionProyectoAssessment' que cumplan dichas condiciones.
                    context["assess"] = AsociacionProyectoAssessment.objects.filter(
                        proyecto=Proyecto.objects.get(codigo=request.session.get('proyectoSeleccionado')),
                        assessment__archivado=0)
                    # Se guardan todos los marcos de seguridad en el contexto.
                    context["marcos"] = AsociacionMarcos.objects.all()

                    # Se renderiza el template en base del contexto.
                    return render(request, self.template_name, context=context)

            # Si tanto el nombre como el selector de marcos están vacíos.
            else:
                # Se crea mensage de error
                messages.error(request, 'Necesitas introducir un nombre para el Assessment')

                '''Inicialización del contexto '''
                # Se inicializa un diccionario llamado context con algunos datos de contexto.
                context = super(assessmentselect, self).get_context_data(**knwargs)
                # Conseguimos los proyectos asociados al usuario.
                context["proyectos"] = AsociacionUsuariosProyecto.objects.filter(usuario=self.request.user)
                # Guardamos en el contexto el proyecto seleccionado, el cual fue guardado anteriormente.
                context["proyectoSelec"] = request.session.get('proyectoSeleccionado')
                context["proyectoSeleccionado"] = True
                # Obtenemos todos los filas de la tabla 'AsociacionProyectoAssessment' que cumplan dichas condiciones.
                context["assess"] = AsociacionProyectoAssessment.objects.filter(
                    proyecto=Proyecto.objects.get(codigo=request.session.get('proyectoSeleccionado')),
                    assessment__archivado=0)
                # Se guardan todos los marcos de seguridad en el contexto.
                context["marcos"] = AsociacionMarcos.objects.all()

                # Se renderiza el template en base del contexto.
                return render(request, self.template_name, context=context)


# Clase para la pagina de Exportaciones
class Exportaciones(LoginRequiredMixin, TemplateView):
    login_url = ""
    redirect_field_name = "redirect_to"
    template_name = "homepage/Exportaciones.html"
    conn = mysql.connector.connect(user='root', password="NTTCSCF2023", host='127.0.0.1', database='nttcs_cf',
                                   auth_plugin='mysql_native_password')  # constante para la conexion con la base de datos

    def get_context_data(self, **knwargs):
        context = super(Exportaciones, self).get_context_data(**knwargs)
        context["proyectos"] = AsociacionUsuariosProyecto.objects.filter(usuario=self.request.user)
        return context

    def PrepararExportacion(self, seleccion, selector):
        ass = Assessmentguardados.objects.get(id_assessment=selector)
        consulta = AssessmentCreados.objects.filter(assessment=ass)
        valores = []
        for i in range(0, len(seleccion)):
            if seleccion[i] == '':
                seleccion = seleccion[0:i]
                break

        for fila in consulta:  # Rellenamos tanto las casillas de respuesta y valoracion
            evgen = AsociacionEvidenciasGenericas.objects.filter(assessment=fila)
            evcre = AsociacionEvidenciasCreadas.objects.filter(id_assessment=fila)
            evidencias = ''
            iniciativas = ''
            for i in evgen:
                if ass.idioma == 'en':
                    evidencias += i.evidencia.evidence_request_references + '\n'
                else:
                    evidencias += i.evidencia_id_es.evidence_request_references + '\n'
                if i.iniciativa != None:
                    iniciativas += i.iniciativa.nombre + '\n'
            for i in evcre:
                evidencias += i.id_evidencia.evidencia_id + '\n'
                if i.iniciativa != None:
                    iniciativas += i.iniciativa.nombre + '\n'

            valor = []

            if "Identificador Control" in seleccion:
                valor += [('Identificador Control', fila.control_id)]
            if "Nombre Control" in seleccion:
                valor += [('Nombre Control', fila.control_name)]
            if "Descripcion Control" in seleccion:
                valor += [('Descripcion Control', fila.descripcion)]
            if "Pregunta" in seleccion:
                valor += [('Pregunta', fila.pregunta)]
            if "Respuesta" in seleccion:
                if fila.respuesta == None:
                    valor += [('Respuesta', fila.respuesta)]
                else:
                    valor += [('Respuesta', BeautifulSoup(fila.respuesta, "lxml").text)]

            if "Valoracion" in seleccion:
                valor += [('Valoracion', fila.valoracion)]
            if "Valoracion Objetivo" in seleccion:
                valor += [('Valoracion Objetivo', fila.valoracionobjetivo)]
            if "Evidencias" in seleccion:
                valor += [('Evidencias', evidencias)]
            if "Iniciativas" in seleccion:
                valor += [('Iniciativas', iniciativas)]

            valores += dict(valor),

        titulos = []
        for i in seleccion:
            titulos += [i]
        now = datetime.now()
        filename = 'Exportaciones/' + selector + '_Export_' + str(now.day) + '_' + str(now.month) + '_' + str(
            now.year) + '_' + str(now.hour) + '_' + str(now.minute)
        return filename, titulos, valores
    # funcion post que recoge los summit del formulario de la pagina.
    def post(self, request, **knwargs):
        selectorProyecto = request.POST.get('selectorProyecto')  # valor de selector de proyecto
        selector = request.POST.get('selector1')
        excel = request.POST.get('excel')
        csvinput = request.POST.get('csv')
        word = request.POST.get('word')
        if 'selectorProyecto' in request.POST:
            context = super(Exportaciones, self).get_context_data(**knwargs)
            context["proyectos"] = AsociacionUsuariosProyecto.objects.filter(usuario=self.request.user)
            context["proyectoSelec"] = selectorProyecto
            context["proyectoSeleccionado"] = True
            request.session["proyectoSeleccionado"] = selectorProyecto
            context["assess"] = AsociacionProyectoAssessment.objects.filter(
                proyecto=Proyecto.objects.get(codigo=selectorProyecto), assessment__archivado=0)
            context["marcos"] = AsociacionMarcos.objects.all()
            return render(request, self.template_name, context=context)
        elif "selector1" in request.POST:
            if selector != 'None':
                seleccion = request.POST.getlist("selector2")
                filename, titulos, valores = self.PrepararExportacion(seleccion,selector)

                if 'csv' == csvinput:
                    filename += '.csv'
                    with open(filename, mode='w', encoding="cp437", errors="replace") as file:
                        writer = csv.DictWriter(file, delimiter=',', fieldnames=titulos)
                        writer.writeheader()

                        for valor in valores:
                            writer.writerow(valor)

                    path = open(filename, 'r')
                    mime_type, _ = mimetypes.guess_type(filename)
                    response = HttpResponse(path, content_type=mime_type)
                    response['Content-Disposition'] = f"attachment; filename={filename}"
                    return response

                if 'excel' == excel:
                    filename += '.xlsx'

                    df = pd.DataFrame(data=valores)
                    writer = pd.ExcelWriter(filename, engine='xlsxwriter')
                    df.to_excel(writer, sheet_name='Sheet1', startrow=1, header=False, index=False)
                    workbook = writer.book
                    worksheet = writer.sheets['Sheet1']
                    column_settings = [{'header': column} for column in df.columns]
                    (max_row, max_col) = df.shape
                    worksheet.add_table(0, 0, max_row, max_col - 1, {'columns': column_settings})
                    worksheet.autofilter(0, 0, max_row, max_col - 1)
                    cell_format = workbook.add_format()
                    cell_format.set_text_wrap()
                    worksheet.set_column('A:I', 20, cell_format)
                    writer.close()

                    with open(filename, "rb") as file:
                        response = HttpResponse(file.read(),
                                                content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
                        response['Content-Disposition'] = f"attachment; filename={filename}"
                    return response

                if 'word' == word:

                    context = super(Exportaciones, self).get_context_data(**knwargs)
                    context["proyectos"] = AsociacionUsuariosProyecto.objects.filter(usuario=self.request.user)
                    context["proyectoSelec"] = request.session["proyectoSeleccionado"]
                    context["proyectoSeleccionado"] = True
                    context["word"] = True
                    request.session["titulos"] = seleccion
                    request.session["selector"] = selector
                    lista = seleccion
                    for i in range(0, 36-len(lista)):
                        lista += ['']
                    context["seleccion"] = lista
                    context["assess"] = AsociacionProyectoAssessment.objects.filter(
                        proyecto=Proyecto.objects.get(codigo=request.session["proyectoSeleccionado"]), assessment__archivado=0)
                    context["marcos"] = AsociacionMarcos.objects.all()


                    return render(request, self.template_name, context=context)
        elif 'lista' in request.POST:
            orden = request.POST.getlist('lista')
            filename, titulos, valores = self.PrepararExportacion(request.session["titulos"], request.session["selector"])

            for i in range(len(orden)-1, 0, -1):
                if orden[i] != '':
                    orden = orden[0:i+1]
                    break
            orden = [orden[i:i + 4] for i in range(0, len(orden), 4)]

            o = []
            for i in orden:
                p = []
                for h in i:
                    if h != "":
                        p += [h]
                o += [p]
            orden = o

            maxLen = 0
            for i in orden:
                if len(i) > maxLen:
                    maxLen = len(i)

            filename += '.docx'
            # create document object
            document = Document()

            for i in valores:
                document.add_heading(i['Identificador Control'] + ", " + i['Nombre Control'], level=1)
                table = document.add_table(rows=0, cols=maxLen*2)
                table.style = 'TableGrid'
                for j in orden:

                    row_cells = table.add_row().cells
                    contador = 0
                    for p in j:
                        if p == '':
                            row_cells[(contador*2)].merge(row_cells[(contador*2)+1])

                        else:
                            row_cells[(contador*2)].text = p

                            if i[str(p)] != None:
                                row_cells[(contador*2)+1].text = i[p]
                            else:
                                row_cells[(contador*2)+1].text = ""
                        contador += 1
                p = document.add_paragraph('')
            # save document
            document.save(filename)
        context = super(Exportaciones, self).get_context_data(**knwargs)
        context["proyectos"] = AsociacionUsuariosProyecto.objects.filter(usuario=self.request.user)
        return render(request, self.template_name, context=context)


# Clase para la pagina de informes
class informes(LoginRequiredMixin, TemplateView):
    login_url = ""
    redirect_field_name = "redirect_to"
    template_name = "homepage/informes.html"
    conn = mysql.connector.connect(user='root', password="NTTCSCF2023", host='127.0.0.1', database='nttcs_cf',
                                   auth_plugin='mysql_native_password')  # constante para la conexion con la base de datos

    def get_context_data(self, **knwargs):
        context = super(informes, self).get_context_data(**knwargs)
        context["assess"] = Assessmentguardados.objects.all()
        return context


# Clase para la pagina de Mantenimiento
class Mantenimiento(LoginRequiredMixin, TemplateView):
    login_url = ""
    redirect_field_name = "redirect_to"
    template_name = "homepage/Mantenimiento.html"


# Clase para la pagina de MantenimientoNivelMadurez
class MantenimientoNivelMadurez(LoginRequiredMixin, TemplateView):
    login_url = ""
    redirect_field_name = "redirect_to"
    template_name = "homepage/MantenimientoNivelMadurez.html"

    # funcion post que recoge los summit del formulario de la pagina.
    def post(self, request, **knwargs):

        if 'insertar' in request.POST:
            Ccmmcod = request.POST.get('Ccmmcod')  # valor del input de ccmmcod
            Description = request.POST.get('Description')  # valor del input de descripcion
            Sublevels = request.POST.get('Sublevels')  # valor del input de sublevels
            if request.POST.get('Percentage') == '':
                Percentage = request.POST.get('Percentage')  # valor del input de percentaje
            else:
                Percentage = float(request.POST.get('Percentage').replace(',', '.'))  # valor del input de percentaje
            if not MaturirtyTable.objects.filter(sublevels=Sublevels).exists():
                if Ccmmcod != '' and Description != '' and Sublevels != '' and Percentage != '':
                    insert = MaturirtyTable(ccmmcod=Ccmmcod, description=Description, sublevels=Sublevels,
                                            percentage=Percentage)  # creamos un nuevo input en la tabla
                    insert.save()
                else:
                    messages.error(request,
                                   'ERROR, debe introducir todos los valores para insertar un nivel de madurez')
            else:
                messages.error(request, 'ERROR, el sublevel debe ser distinto a uno ya existente')
        elif 'modificar' in request.POST:
            Ccmmcod = request.POST.get('Ccmmcod')  # valor del input de ccmmcod
            Description = request.POST.get('Description')  # valor del input de descripcion
            Sublevels = request.POST.get('Sublevels')  # valor del input de sublevels
            Percentage = float(request.POST.get('Percentage').replace(',', '.'))  # valor del input de percentaje
            consulta = MaturirtyTable.objects.get(
                sublevels=Sublevels)  # si esta en la tabla seleccionamos el ojeto en la tabla
            consulta.ccmmcod = Ccmmcod
            consulta.description = Description
            consulta.percentage = Percentage
            consulta.save()  # fijamos los valores y los guardamos.
        elif 'eliminar' in request.POST:
            Sublevels = request.POST.get('Sublevels')  # valor del input de sublevels
            consulta = MaturirtyTable.objects.get(sublevels=Sublevels)
            consulta.delete()  # seleccionamos el objeto de la ultima busqueda y lo eliminamos.

        page = self.request.GET.get('page', 1)
        paginator = Paginator(MaturirtyTable.objects.all(), 50)
        context = super(MantenimientoNivelMadurez, self).get_context_data(**knwargs)
        context["entity"] = paginator.page(page)
        context["paginator"] = paginator
        context["lenConsulta"] = len(MaturirtyTable.objects.all())
        return render(request, self.template_name,
                      context=context)  # siempre retornamos el valor con la tabla completa.

    # funcion que envia el contexto de la pagina.
    def get_context_data(self, **knwargs):
        page = self.request.GET.get('page', 1)
        paginator = Paginator(MaturirtyTable.objects.all(), 50)
        context = super(MantenimientoNivelMadurez, self).get_context_data(**knwargs)
        context["entity"] = paginator.page(page)
        context["paginator"] = paginator
        context["lenConsulta"] = len(MaturirtyTable.objects.all())
        return context


# Clase para la pagina de inicio de sesion


# Clase para la pagina de menu
class menu(LoginRequiredMixin, TemplateView):
    login_url = ""
    redirect_field_name = "redirect_to"
    template_name = "homepage/menu.html"


# Clase para la pagina de MantenimientoDominios
class MantenimientoDominios(LoginRequiredMixin, TemplateView):
    login_url = ""
    redirect_field_name = "redirect_to"
    template_name = "homepage/MantenimientoDominios.html"

    # funcion post que recoge los summit del formulario de la pagina.
    def post(self, request, **knwargs):

        if 'insertar' in request.POST:
            identifier = request.POST.get('identifier')  # valor del input de identifier
            domain = request.POST.get('domain')  # valor del input de domain
            security_privacy_by_design_s_p_principles = request.POST.get(
                'security_privacy_by_design_s_p_principles')  # valor del input de security_privacy_by_design_s_p_principles
            principle_intent = request.POST.get('principle_intent')  # valor del input de principle_intent
            if not Domains.objects.filter(identifier=identifier).exists():
                if identifier != '' and domain != '' and security_privacy_by_design_s_p_principles != '' and principle_intent != '':
                    insert = Domains(identifier=identifier, domain=domain,
                                     security_privacy_by_design_s_p_principles=security_privacy_by_design_s_p_principles,
                                     principle_intent=principle_intent, )  # creamos un nuevo input en la tabla
                    insert.save()
                else:
                    messages.error(request, 'ERROR, debe introducir todos los valores para insertar un dominio')
            else:
                messages.error(request, 'ERROR, el identifier debe ser distinto a uno ya existente')
        elif 'modificar' in request.POST:
            identifier = request.POST.get('identifier')  # valor del input de identifier
            domain = request.POST.get('domain')  # valor del input de domain
            security_privacy_by_design_s_p_principles = request.POST.get(
                'security_privacy_by_design_s_p_principles')  # valor del input de security_privacy_by_design_s_p_principles
            principle_intent = request.POST.get('principle_intent')  # valor del input de principle_intent
            consulta = Domains.objects.get(
                identifier=identifier)  # consulta para seleccionar el objeto que corresponde con la ultima busqueda
            consulta.domain = domain
            consulta.security_privacy_by_design_s_p_principles = security_privacy_by_design_s_p_principles
            consulta.principle_intent = principle_intent
            consulta.save()  # fijamos los valores y los guardamos.
        elif 'eliminar' in request.POST:
            identifier = request.POST.get('identifier')  # valor del input de identifier

            consulta = Domains.objects.get(identifier=identifier)
            consulta.delete()  # seleccionamos el objeto de la ultima busqueda y lo eliminamos.

        page = self.request.GET.get('page', 1)
        paginator = Paginator(Domains.objects.all(), 50)
        context = super(MantenimientoDominios, self).get_context_data(**knwargs)
        context["entity"] = paginator.page(page)
        context["paginator"] = paginator
        context["lenConsulta"] = len(Domains.objects.all())
        return render(request, self.template_name,
                      context=context)  # siempre retornamos el valor con la tabla completa.

    # funcion que envia el contexto de la pagina.
    def get_context_data(self, **knwargs):
        page = self.request.GET.get('page', 1)
        paginator = Paginator(Domains.objects.all(), 50)
        context = super(MantenimientoDominios, self).get_context_data(**knwargs)
        context["entity"] = paginator.page(page)
        context["paginator"] = paginator
        context["lenConsulta"] = len(Domains.objects.all())
        return context

    # Funcion utilizada para eliminar el valor seleccionado de la tabla


# Clase para la pagina de MantenimientoEvidencias
class MantenimientoEvidencias(LoginRequiredMixin, TemplateView):
    login_url = ""
    redirect_field_name = "redirect_to"
    template_name = "homepage/MantenimientoEvidencias.html"

    # funcion post que recoge los summit del formulario de la pagina.
    def post(self, request, **knwargs):

        if 'insertar' in request.POST:

            evidence_request_references = request.POST.get(
                'evidence_request_references')  # valor del input de evidence_request_references
            area_of_focus = request.POST.get('area_of_focus')  # valor del input de area_of_focus
            artifact = request.POST.get('artifact')  # valor del input de artifact
            artifact_description = request.POST.get('artifact_description')  # valor del input de artifact_description
            control_mappings = request.POST.get('control_mappings')  # valor del input de control_mappings
            if not Evidencerequestcatalog.objects.filter(
                    evidence_request_references=evidence_request_references).exists():
                if evidence_request_references != '' and area_of_focus != '' and artifact != '' and artifact_description != '' and control_mappings != '':
                    insert = Evidencerequestcatalog(evidence_request_references=evidence_request_references,
                                                    area_of_focus=area_of_focus, artifact=artifact,
                                                    artifact_description=artifact_description,
                                                    control_mappings=control_mappings)  # creamos un nuevo input en la tabla
                    insert.save()
                else:
                    messages.error(request, 'ERROR, debe introducir todos los valores para insertar una Evidencia')
            else:
                messages.error(request, 'ERROR, la evidence_request_references debe ser distinto a uno ya existente')
        elif 'modificar' in request.POST:
            evidence_request_references = request.POST.get(
                'evidence_request_references')  # valor del input de evidence_request_references
            area_of_focus = request.POST.get('area_of_focus')  # valor del input de area_of_focus
            artifact = request.POST.get('artifact')  # valor del input de artifact
            artifact_description = request.POST.get('artifact_description')  # valor del input de artifact_description
            control_mappings = request.POST.get('control_mappings')  # valor del input de control_mappings
            consulta = Evidencerequestcatalog.objects.get(
                evidence_request_references=evidence_request_references)  # consulta para seleccionar el objeto que corresponde con la ultima busqueda
            consulta.area_of_focus = area_of_focus
            consulta.artifact = artifact
            consulta.artifact_description = artifact_description
            consulta.control_mappings = control_mappings
            consulta.save()  # fijamos los valores y los guardamos.
        elif 'eliminar' in request.POST:
            evidence_request_references = request.POST.get(
                'evidence_request_references')  # valor del input de evidence_request_references

            consulta = Evidencerequestcatalog.objects.get(evidence_request_references=evidence_request_references)
            consulta.delete()  # seleccionamos el objeto de la ultima busqueda y lo eliminamos.

        page = self.request.GET.get('page', 1)
        paginator = Paginator(Evidencerequestcatalog.objects.all(), 50)
        context = super(MantenimientoEvidencias, self).get_context_data(**knwargs)
        context["entity"] = paginator.page(page)
        context["paginator"] = paginator
        context["lenConsulta"] = len(Evidencerequestcatalog.objects.all())
        return render(request, self.template_name,
                      context=context)  # siempre retornamos el valor con la tabla completa.

    # funcion que envia el contexto de la pagina.
    def get_context_data(self, **knwargs):
        page = self.request.GET.get('page', 1)
        paginator = Paginator(Evidencerequestcatalog.objects.all(), 50)
        context = super(MantenimientoEvidencias, self).get_context_data(**knwargs)
        context["entity"] = paginator.page(page)
        context["paginator"] = paginator
        context["lenConsulta"] = len(Evidencerequestcatalog.objects.all())
        return context


# Clase para la pagina de MantenimientoEvidencias
class MantenimientoEvidenciasEs(LoginRequiredMixin, TemplateView):
    login_url = ""
    redirect_field_name = "redirect_to"
    template_name = "homepage/MantenimientoEvidenciasEs.html"

    # funcion post que recoge los summit del formulario de la pagina.
    def post(self, request, **knwargs):

        if 'insertar' in request.POST:

            evidence_request_references = request.POST.get(
                'evidence_request_references')  # valor del input de evidence_request_references
            area_of_focus = request.POST.get('area_of_focus')  # valor del input de area_of_focus
            artifact = request.POST.get('artifact')  # valor del input de artifact
            artifact_description = request.POST.get('artifact_description')  # valor del input de artifact_description
            control_mappings = request.POST.get('control_mappings')  # valor del input de control_mappings
            if not EvidencerequestcatalogEs.objects.filter(
                    evidence_request_references=evidence_request_references).exists():
                if evidence_request_references != '' and area_of_focus != '' and artifact != '' and artifact_description != '' and control_mappings != '':
                    insert = EvidencerequestcatalogEs(evidence_request_references=evidence_request_references,
                                                      area_of_focus=area_of_focus, artifact=artifact,
                                                      artifact_description=artifact_description,
                                                      control_mappings=control_mappings)  # creamos un nuevo input en la tabla
                    insert.save()
                else:
                    messages.error(request, 'ERROR, debe introducir todos los valores para insertar una Evidencia')
            else:
                messages.error(request, 'ERROR, la evidence_request_references debe ser distinto a uno ya existente')
        elif 'modificar' in request.POST:
            evidence_request_references = request.POST.get(
                'evidence_request_references')  # valor del input de evidence_request_references
            area_of_focus = request.POST.get('area_of_focus')  # valor del input de area_of_focus
            artifact = request.POST.get('artifact')  # valor del input de artifact
            artifact_description = request.POST.get('artifact_description')  # valor del input de artifact_description
            control_mappings = request.POST.get('control_mappings')  # valor del input de control_mappings
            consulta = EvidencerequestcatalogEs.objects.get(
                evidence_request_references=evidence_request_references)  # consulta para seleccionar el objeto que corresponde con la ultima busqueda
            consulta.area_of_focus = area_of_focus
            consulta.artifact = artifact
            consulta.artifact_description = artifact_description
            consulta.control_mappings = control_mappings
            consulta.save()  # fijamos los valores y los guardamos.
        elif 'eliminar' in request.POST:
            evidence_request_references = request.POST.get(
                'evidence_request_references')  # valor del input de evidence_request_references

            consulta = EvidencerequestcatalogEs.objects.get(evidence_request_references=evidence_request_references)
            consulta.delete()  # seleccionamos el objeto de la ultima busqueda y lo eliminamos.

        page = self.request.GET.get('page', 1)
        paginator = Paginator(EvidencerequestcatalogEs.objects.all(), 50)
        context = super(MantenimientoEvidenciasEs, self).get_context_data(**knwargs)
        context["entity"] = paginator.page(page)
        context["paginator"] = paginator
        context["lenConsulta"] = len(EvidencerequestcatalogEs.objects.all())
        return render(request, self.template_name,
                      context=context)  # siempre retornamos el valor con la tabla completa.

    # funcion que envia el contexto de la pagina.
    def get_context_data(self, **knwargs):
        page = self.request.GET.get('page', 1)
        paginator = Paginator(EvidencerequestcatalogEs.objects.all(), 50)
        context = super(MantenimientoEvidenciasEs, self).get_context_data(**knwargs)
        context["entity"] = paginator.page(page)
        context["paginator"] = paginator
        context["lenConsulta"] = len(EvidencerequestcatalogEs.objects.all())
        return context


# Clase para la pagina de MantenimientoPreguntas
class MantenimientoPreguntas(LoginRequiredMixin, TemplateView):
    login_url = ""
    redirect_field_name = "redirect_to"
    template_name = "homepage/MantenimientoPreguntas.html"

    # funcion post que recoge los summit del formulario de la pagina.
    def post(self, request, **knwargs):

        if 'insertar' in request.POST:
            control_question = request.POST.get('control_question')  # valor del input de control_question
            control_description = request.POST.get('control_description')  # valor del input de control_description
            id = request.POST.get('id')  # valor del input de id
            if not Assessment.objects.filter(id=id).exists():
                if control_question != '' and control_description != '' and id != '':
                    insert = Assessment(id=id, control_question=control_question,
                                        control_description=control_description)  # creamos un nuevo input en la tabla
                    insert.save()
                else:
                    messages.error(request, 'ERROR, debe introducir todos los valores para insertar una Pregunta')
            else:
                messages.error(request, 'ERROR, el id debe ser distinto a uno ya existente')
        elif 'modificar' in request.POST:
            control_question = request.POST.get('control_question')  # valor del input de control_question
            control_description = request.POST.get('control_description')  # valor del input de control_description
            id = request.POST.get('id')  # valor del input de id
            consulta = Assessment.objects.get(id=id)  # si esta en la tabla seleccionamos el ojeto en la tabla
            consulta.control_description = control_description
            consulta.control_question = control_question
            consulta.save()  # fijamos los valores y los guardamos..
        elif 'eliminar' in request.POST:
            id = request.POST.get('id')  # valor del input de id

            consulta = Assessment.objects.get(id=id)
            consulta.delete()  # seleccionamos el objeto de la ultima busqueda y lo eliminamos.

        page = self.request.GET.get('page', 1)
        paginator = Paginator(Assessment.objects.all(), 50)
        context = super(MantenimientoPreguntas, self).get_context_data(**knwargs)
        context["entity"] = paginator.page(page)
        context["paginator"] = paginator
        context["lenConsulta"] = len(Assessment.objects.all())
        return render(request, self.template_name,
                      context=context)  # siempre retornamos el valor con la tabla completa.

    # funcion que envia el contexto de la pagina.
    def get_context_data(self, **knwargs):
        page = self.request.GET.get('page', 1)
        paginator = Paginator(Assessment.objects.all(), 50)
        context = super(MantenimientoPreguntas, self).get_context_data(**knwargs)
        context["entity"] = paginator.page(page)
        context["paginator"] = paginator
        context["lenConsulta"] = len(Assessment.objects.all())
        return context


class MantenimientoPreguntasEs(LoginRequiredMixin, TemplateView):
    login_url = ""
    redirect_field_name = "redirect_to"
    template_name = "homepage/MantenimientoPreguntasEs.html"

    # funcion post que recoge los summit del formulario de la pagina.
    def post(self, request, **knwargs):

        if 'insertar' in request.POST:
            control_question = request.POST.get('control_question')  # valor del input de control_question
            control_description = request.POST.get('control_description')  # valor del input de control_description
            id = request.POST.get('id')  # valor del input de id
            if not AssessmentEs.objects.filter(id=id).exists():
                if control_question != '' and control_description != '' and id != '':
                    insert = AssessmentEs(id=id, control_question=control_question,
                                          control_description=control_description)  # creamos un nuevo input en la tabla
                    insert.save()
                else:
                    messages.error(request, 'ERROR, debe introducir todos los valores para insertar una Pregunta')
            else:
                messages.error(request, 'ERROR, el id debe ser distinto a uno ya existente')
        elif 'modificar' in request.POST:
            control_question = request.POST.get('control_question')  # valor del input de control_question
            control_description = request.POST.get('control_description')  # valor del input de control_description
            id = request.POST.get('id')  # valor del input de id
            consulta = AssessmentEs.objects.get(id=id)  # si esta en la tabla seleccionamos el ojeto en la tabla
            consulta.control_description = control_description
            consulta.control_question = control_question
            consulta.save()  # fijamos los valores y los guardamos..
        elif 'eliminar' in request.POST:
            id = request.POST.get('id')  # valor del input de id

            consulta = AssessmentEs.objects.get(id=id)
            consulta.delete()  # seleccionamos el objeto de la ultima busqueda y lo eliminamos.

        page = self.request.GET.get('page', 1)
        paginator = Paginator(AssessmentEs.objects.all(), 50)
        context = super(MantenimientoPreguntasEs, self).get_context_data(**knwargs)
        context["entity"] = paginator.page(page)
        context["paginator"] = paginator
        context["lenConsulta"] = len(AssessmentEs.objects.all())
        return render(request, self.template_name,
                      context=context)  # siempre retornamos el valor con la tabla completa.

    # funcion que envia el contexto de la pagina.
    def get_context_data(self, **knwargs):
        page = self.request.GET.get('page', 1)
        paginator = Paginator(AssessmentEs.objects.all(), 50)
        context = super(MantenimientoPreguntasEs, self).get_context_data(**knwargs)
        context["entity"] = paginator.page(page)
        context["paginator"] = paginator
        context["lenConsulta"] = len(AssessmentEs.objects.all())
        return context


# Clase para la pagina de MantenimientoMarcosExistentes
class MantenimientoMarcosExistentes(LoginRequiredMixin, TemplateView):
    login_url = ""
    redirect_field_name = "redirect_to"
    template_name = "homepage/MantenimientoMarcosExistentes.html"

    # funcion post que recoge los summit del formulario de la pagina.
    def post(self, request, **knwargs):
        if 'insertar' in request.POST:
            marco_id = request.POST.get('marco_id')  # valor del input de marco_id
            nombre_tabla = request.POST.get('nombre_tabla')  # valor del input de nombre_tabla
            if not AsociacionMarcos.objects.filter(marco_id=marco_id).exists():
                if marco_id != '' and nombre_tabla != '':
                    insert = AsociacionMarcos(marco_id=marco_id,
                                              nombre_tabla=nombre_tabla)  # creamos un nuevo input en la tabla
                    insert.save()
                else:
                    messages.error(request, 'ERROR, debe introducir todos los valores para insertar un marco')
            else:
                messages.error(request, 'ERROR, el marco_id debe ser distinto a uno ya existente')
        elif 'modificar' in request.POST:
            marco_id = request.POST.get('marco_id')  # valor del input de marco_id
            nombre_tabla = request.POST.get('nombre_tabla')  # valor del input de nombre_tabla

            consulta = AsociacionMarcos.objects.get(
                marco_id=marco_id)  # si esta en la tabla seleccionamos el ojeto en la tabla
            consulta.nombre_tabla = nombre_tabla
            consulta.save()  # fijamos los valores y los guardamos.
        elif 'eliminar' in request.POST:
            marco_id = request.POST.get('marco_id')  # valor del input de marco_id
            consulta = AsociacionMarcos.objects.get(marco_id=marco_id)
            consulta.delete()  # seleccionamos el objeto de la ultima busqueda y lo eliminamos.

        page = self.request.GET.get('page', 1)
        paginator = Paginator(AsociacionMarcos.objects.all(), 50)
        context = super(MantenimientoMarcosExistentes, self).get_context_data(**knwargs)
        context["entity"] = paginator.page(page)
        context["paginator"] = paginator
        context["lenConsulta"] = len(AsociacionMarcos.objects.all())
        return render(request, self.template_name,
                      context=context)  # siempre retornamos el valor con la tabla completa.

    # funcion que envia el contexto de la pagina.
    def get_context_data(self, **knwargs):
        page = self.request.GET.get('page', 1)
        paginator = Paginator(AsociacionMarcos.objects.all(), 50)
        context = super(MantenimientoMarcosExistentes, self).get_context_data(**knwargs)
        context["entity"] = paginator.page(page)
        context["paginator"] = paginator
        context["lenConsulta"] = len(AsociacionMarcos.objects.all())
        return context

    # Funcion utilizada para eliminar el valor seleccionado de la tabla


class listadoControles(LoginRequiredMixin, ListView):
    model = NttcsCf20231

    def get_queryset(self):
        return self.model.objects.all().values()

    def get(self, request, *args, **kwargs):
        inicio = int(request.GET.get('inicio'))
        fin = int(request.GET.get('limite'))
        tiempo_inicial = time()
        data = self.get_queryset()
        list_data = []
        for indice, valor in enumerate(data[inicio: inicio + fin]):
            list_data.append(valor)
        tiempo_final = time() - tiempo_inicial
        print(f'Tiempo de Ejecución: {tiempo_final}')

        data = {
            'length': data.count(),
            'object': list_data
        }

        return HttpResponse(json.dumps(data), 'application/json')


# Clase para la pagina de MantenimientoControlesNTTCS
class MantenimientoControlesNTTCS(LoginRequiredMixin, TemplateView):
    login_url = ""
    redirect_field_name = "redirect_to"
    template_name = "homepage/MantenimientoControlesNTTCS.html"

    # funcion post que recoge los summit del formulario de la pagina.
    def post(self, request, **knwargs):

        if 'insertar' in request.POST:
            domain = request.POST.get('domain')  # valor del input de domain
            selected_y_n_field = request.POST.get('selected_y_n_field')  # valor del input de selected_y_n_field
            control = request.POST.get('control')  # valor del input de control
            id = request.POST.get('id')  # valor del input de id
            control_description = request.POST.get('control_description')  # valor del input de control_description
            relative_control_weighting = request.POST.get(
                'relative_control_weighting')  # valor del input de relative_control_weighting
            function_grouping = request.POST.get('function_grouping')  # valor del input de function_grouping
            assesed_result = request.POST.get('assesed_result')  # valor del input de assesed_result
            numeric_result = request.POST.get('numeric_result')  # valor del input de numeric_result
            weighted_numeric_result = request.POST.get(
                'weighted_numeric_result')  # valor del input de weighted_numeric_result
            assessment_comments = request.POST.get('assessment_comments')  # valor del input de assessment_comments
            relative_result_by_function = request.POST.get(
                'relative_result_by_function')  # valor del input de relative_result_by_function
            relative_result_by_domain = request.POST.get(
                'relative_result_by_domain')  # valor del input de relative_result_by_domain
            if not NttcsCf20231.objects.filter(id=id).exists():
                if domain != '' and selected_y_n_field != '' and control != '' and id != '' and control_description != '' and relative_control_weighting != '' and function_grouping != '' and assesed_result != '' and numeric_result != '' and weighted_numeric_result != '' and assessment_comments != '' and relative_result_by_function != '' and relative_result_by_domain != '':
                    insert = NttcsCf20231(domain=domain, selected_y_n_field=selected_y_n_field, id=id, control=control,
                                          control_description=control_description,
                                          relative_control_weighting=relative_control_weighting,
                                          function_grouping=function_grouping, assesed_result=assesed_result,
                                          numeric_result=numeric_result,
                                          weighted_numeric_result=weighted_numeric_result,
                                          assessment_comments=assessment_comments,
                                          relative_result_by_function=relative_result_by_function,
                                          relative_result_by_domain=relative_result_by_domain)  # creamos un nuevo input en la tabla
                    insert.save()
                else:
                    messages.error(request, 'ERROR, debe introducir todos los valores para insertar un control')
            else:
                messages.error(request, 'ERROR, el id debe ser distinto a uno ya existente')
        elif 'modificar' in request.POST:
            domain = request.POST.get('domain')  # valor del input de domain
            selected_y_n_field = request.POST.get('selected_y_n_field')  # valor del input de selected_y_n_field
            control = request.POST.get('control')  # valor del input de control
            id = request.POST.get('id')  # valor del input de id
            control_description = request.POST.get('control_description')  # valor del input de control_description
            relative_control_weighting = request.POST.get(
                'relative_control_weighting')  # valor del input de relative_control_weighting
            function_grouping = request.POST.get('function_grouping')  # valor del input de function_grouping
            assesed_result = request.POST.get('assesed_result')  # valor del input de assesed_result
            numeric_result = request.POST.get('numeric_result')  # valor del input de numeric_result
            weighted_numeric_result = request.POST.get(
                'weighted_numeric_result')  # valor del input de weighted_numeric_result
            assessment_comments = request.POST.get('assessment_comments')  # valor del input de assessment_comments
            relative_result_by_function = request.POST.get(
                'relative_result_by_function')  # valor del input de relative_result_by_function
            relative_result_by_domain = request.POST.get(
                'relative_result_by_domain')  # valor del input de relative_result_by_domain
            consulta = NttcsCf20231.objects.get(id=id)  # si esta en la tabla seleccionamos el ojeto en la tabla
            consulta.domain = domain
            consulta.selected_y_n_field = selected_y_n_field
            consulta.control = control
            consulta.control_description = control_description
            consulta.relative_control_weighting = relative_control_weighting
            consulta.function_grouping = function_grouping
            consulta.assesed_result = assesed_result
            consulta.numeric_result = numeric_result
            consulta.weighted_numeric_result = weighted_numeric_result
            consulta.assessment_comments = assessment_comments
            consulta.relative_result_by_function = relative_result_by_function
            consulta.relative_result_by_domain = relative_result_by_domain
            consulta.save()  # fijamos los valores y los guardamos.
        elif 'eliminar' in request.POST:
            id = request.POST.get('id')  # valor del input de id
            consulta = NttcsCf20231.objects.get(id=id)
            consulta.delete()  # seleccionamos el objeto de la ultima busqueda y lo eliminamos.

        page = self.request.GET.get('page', 1)
        paginator = Paginator(NttcsCf20231.objects.all(), 50)
        context = super(MantenimientoControlesNTTCS, self).get_context_data(**knwargs)
        context["entity"] = paginator.page(page)
        context["paginator"] = paginator
        context["lenConsulta"] = len(NttcsCf20231.objects.all())
        return render(request, self.template_name,
                      context=context)  # siempre retornamos el valor con la tabla completa.

    # funcion que envia el contexto de la pagina.
    def get_context_data(self, **knwargs):

        page = self.request.GET.get('page', 1)
        paginator = Paginator(NttcsCf20231.objects.all(), 50)
        context = super(MantenimientoControlesNTTCS, self).get_context_data(**knwargs)
        context["entity"] = paginator.page(page)
        context["paginator"] = paginator
        context["lenConsulta"] = len(NttcsCf20231.objects.all())
        return context


# Clase para la pagina de MantenimientoMapeoMarcos
class MantenimientoMapeoMarcos(LoginRequiredMixin, TemplateView):
    login_url = ""
    redirect_field_name = "redirect_to"
    template_name = "homepage/MantenimientoMapeoMarcos.html"
    conn = mysql.connector.connect(user='root', password="NTTCSCF2023", host='127.0.0.1', database='nttcs_cf',
                                   auth_plugin='mysql_native_password')

    # funcion que envia el contexto de la pagina.
    def get_context_data(self, **knwargs):
        if self.request.GET.get('page'):
            consulta = MapeoMarcos.objects.all().values_list('ntt_id', AsociacionMarcos.objects.get(
                marco_id=self.request.session["seleccion"]).nombre_tabla.lower())
            page = self.request.GET.get('page', 1)
            paginator = Paginator(consulta, 50)
            context = super(MantenimientoMapeoMarcos, self).get_context_data(**knwargs)
            context["assess"] = AsociacionMarcos.objects.all()
            context["entity"] = paginator.page(page)
            context["paginator"] = paginator  # fijamos la tabla a el valor seleccionado
            context["seleccionado"] = True
            context['marcoSeleccionado'] = self.request.session["seleccion"]
        else:
            context = super(MantenimientoMapeoMarcos, self).get_context_data(**knwargs)
            context["assess"] = AsociacionMarcos.objects.all()

            context["lenConsulta"] = 1
            context["seleccionado"] = False
        return context

    # funcion que envia el contexto de la pagina.
    def post(self, request, **knwargs):

        if 'selector' in request.POST:  # if que recoge la pulsacion del boton de seleccion
            selector = request.POST.get('selector')  # guardamos el valor del selecctor de marcos
            if selector == 'None':
                context = super(MantenimientoMapeoMarcos, self).get_context_data(**knwargs)
                context["assess"] = AsociacionMarcos.objects.all()
                context["lenConsulta"] = 1
                context["seleccionado"] = False
                return render(request, self.template_name, context=context)
            else:

                consulta = MapeoMarcos.objects.all().values_list('ntt_id', AsociacionMarcos.objects.get(
                    marco_id=selector).nombre_tabla.lower())
                page = request.GET.get('page', 1)
                paginator = Paginator(consulta, 50)
                context = super(MantenimientoMapeoMarcos, self).get_context_data(**knwargs)
                context["assess"] = AsociacionMarcos.objects.all()
                context["entity"] = paginator.page(page)
                context["paginator"] = paginator  # fijamos la tabla a el valor seleccionado
                context["lenConsulta"] = 5
                context["seleccionado"] = True
                context['marcoSeleccionado'] = selector
                request.session["seleccion"] = selector  # guardamos la seleecion del marco
                return render(request, self.template_name, context=context)

        elif 'modificar' in request.POST:
            ntt_id = request.POST.get('ntt_id')  # valor del input de ntt_id
            marco = request.POST.get('marco')  # valor del input de marco
            s = AsociacionMarcos.objects.get(marco_id=request.session["seleccion"]).nombre_tabla.lower()

            c = MapeoMarcos.objects.get(ntt_id=ntt_id)
            setattr(c, s, int(marco))
            c.save()


        elif 'eliminar' in request.POST:
            ntt_id = request.POST.get('ntt_id')  # valor del input de ntt_id
            c = MapeoMarcos.objects.get(ntt_id=ntt_id)
            c.delete()



        elif 'insertar' in request.POST:
            ntt_id = request.POST.get('ntt_id')  # valor del input de ntt_id
            marco = request.POST.get('marco')  # valor del input de marco
            if ntt_id != '' and marco != '':
                if not MapeoMarcos.objects.filter(ntt_id=ntt_id).exists():
                    s = AsociacionMarcos.objects.get(marco_id=request.session["seleccion"]).nombre_tabla.lower()
                    c = MapeoMarcos(ntt_id=ntt_id)
                    setattr(c, s, int(marco))
                    c.save()
                else:
                    messages.error(request, 'ERROR,el valor de id ya existe')
            else:
                messages.error(request, 'ERROR, debe introducir todos los valores para insertar un marco')
        elif 'ant' in request.POST:
            if request.session['inicio'] > 0:
                request.session['inicio'] = request.session['inicio'] - 100
        elif 'sig' in request.POST:
            if request.session['inicio'] < AsociacionMarcos.objects.all().count():
                request.session['inicio'] = request.session['inicio'] + 100
        consulta = MapeoMarcos.objects.all().values_list('ntt_id', AsociacionMarcos.objects.get(
            marco_id=request.session["seleccion"]).nombre_tabla.lower())
        page = request.GET.get('page', 1)
        paginator = Paginator(consulta, 50)
        context = super(MantenimientoMapeoMarcos, self).get_context_data(**knwargs)
        context["assess"] = AsociacionMarcos.objects.all()
        context["entity"] = paginator.page(page)
        context["paginator"] = paginator  # fijamos la tabla a el valor seleccionado
        context["seleccionado"] = True
        context['marcoSeleccionado'] = request.session["seleccion"]
        return render(request, self.template_name,
                      context=context)  # siempre retornamos el valor con la tabla completa.


# Clase para la pagina de MantenimientoAssessmentArchivados
class MantenimientoAssessmentArchivados(LoginRequiredMixin, TemplateView):
    login_url = ""
    redirect_field_name = "redirect_to"
    template_name = "homepage/MantenimientoAssessmentArchivados.html"

    # funcion que envia el contexto de la pagina.
    def post(self, request, **knwargs):
        if 'selector' in request.POST:  # if que recoge la pulsacion del boton de seleccion
            selector = request.POST.get('selector')  # guardamos el valor del selecctor de marcos
            if selector == 'None':
                context = super(MantenimientoAssessmentArchivados, self).get_context_data(**knwargs)
                context["assess"] = Assessmentguardados.objects.filter(archivado=1)
                context["seleccionado"] = False
                return render(request, self.template_name, context=context)
            else:
                assGuardado = Assessmentguardados.objects.get(id_assessment=selector)
                # realizamos la consulta para obtener los contrloles del marco seleccionado
                page = self.request.GET.get('page', 1)
                paginator = Paginator(AssessmentCreados.objects.filter(assessment=assGuardado), 50)
                context = super(MantenimientoAssessmentArchivados, self).get_context_data(**knwargs)
                context["assess"] = Assessmentguardados.objects.filter(archivado=1)
                context["entity"] = paginator.page(page)
                context["paginator"] = paginator  # fijamos la tabla a el valor seleccionado
                context["seleccionado"] = True
                context['marcoSeleccionado'] = selector
                request.session["seleccion"] = selector
                request.session["seleccionado"] = True
                return render(request, self.template_name, context=context)

        elif 'modificar' in request.POST:
            id = request.POST.get('id')  # valor del input de id
            nombre = request.POST.get('nombre')  # valor del input de nombre
            descripcion = request.POST.get('descripcion')  # valor del input de descripcion
            Pregunta = request.POST.get('pregunta')  # valor del input de Pregunta
            criterio = request.POST.get('criterio')  # valor del input de criterio
            respuesta = request.POST.get('respuesta')  # valor del input de respuesta
            valoracion = request.POST.get('valoracion')  # valor del input de respuesta
            valoracionobjetivo = request.POST.get('valoracionobjetivo')  # valor del input de evidencia
            c = AssessmentCreados.objects.get(assessment=request.session.get('seleccion'), control_id=id)
            c.control_name = nombre
            c.descripcion = descripcion
            c.pregunta = Pregunta
            c.criterio = criterio
            c.respuesta = respuesta
            c.valoracion = valoracion
            c.valoracionobjetivo = valoracionobjetivo
            c.save()


        elif 'eliminar' in request.POST:
            id = request.POST.get('id')  # valor del input de id
            c = AssessmentCreados.objects.filter(assessment=request.session.get('seleccion'), control_id=id)
            c.delete()  # seleccionamos el objeto de la ultima busqueda y lo eliminamos.

        elif 'insertar' in request.POST:
            id = request.POST.get('id')  # valor del input de id
            nombre = request.POST.get('nombre')  # valor del input de nombre
            descripcion = request.POST.get('descripcion')  # valor del input de descripcion
            Pregunta = request.POST.get('pregunta')  # valor del input de Pregunta
            criterio = request.POST.get('criterio')  # valor del input de criterio
            respuesta = request.POST.get('respuesta')  # valor del input de respuesta
            valoracion = request.POST.get('valoracion')  # valor del input de respuesta
            valoracionobjetivo = request.POST.get('valoracionobjetivo')  # valor del input de evidencia
            if id != '' and descripcion != '' and Pregunta != '' and criterio != '' and respuesta != '' and valoracion != '' and valoracionobjetivo != '':
                if not AssessmentCreados.objects.filter(assessment=request.session.get('seleccion'), control_id=id):
                    a = AssessmentCreados(assessment=request.session.get('seleccion'), control_id=id,
                                          control_name=nombre,
                                          descripcion=descripcion, pregunta=Pregunta,
                                          criteriovaloracion=criterio, valoracionobjetivo=valoracionobjetivo,
                                          respuesta=respuesta, valoracion=valoracion)
                    a.save()
                else:
                    messages.error(request, 'ERROR,el valor de id ya existe')
            else:
                messages.error(request, 'ERROR, debe introducir todos los valores para insertar un marco')

        elif 'eliminarAssessment' in request.POST:
            consulta = Assessmentguardados.objects.get(id_assessment=request.session.get('seleccion'))
            consulta.delete()
            return redirect('menu')

        elif 'desarchivar' in request.POST:
            consulta = Assessmentguardados.objects.get(
                id_assessment=request.session.get('seleccion'))  # colsulta para la selecionar el assesment
            consulta.archivado = 0  # ponemos el valor de archivado a 0
            consulta.save()

        assGuardado = Assessmentguardados.objects.get(id_assessment=request.session["seleccion"])
        # realizamos la consulta para obtener los contrloles del marco seleccionado
        page = self.request.GET.get('page', 1)
        paginator = Paginator(AssessmentCreados.objects.filter(assessment=assGuardado), 50)
        context = super(MantenimientoAssessmentArchivados, self).get_context_data(**knwargs)
        context["assess"] = Assessmentguardados.objects.filter(archivado=1)
        context["entity"] = paginator.page(page)
        context["paginator"] = paginator  # fijamos la tabla a el valor seleccionado
        return render(request, self.template_name,
                      context=context)  # siempre retornamos el valor con la tabla completa.

    # funcion que envia el contexto de la pagina.
    def get_context_data(self, **knwargs):
        if self.request.GET.get('page'):
            assGuardado = Assessmentguardados.objects.get(id_assessment=self.request.session["seleccion"])
            # realizamos la consulta para obtener los contrloles del marco seleccionado
            page = self.request.GET.get('page', 1)
            paginator = Paginator(AssessmentCreados.objects.filter(assessment=assGuardado), 50)
            context = super(MantenimientoAssessmentArchivados, self).get_context_data(**knwargs)
            context["assess"] = Assessmentguardados.objects.filter(archivado=1)
            context["entity"] = paginator.page(page)
            context["paginator"] = paginator  # fijamos la tabla a el valor seleccionado
            context["seleccionado"] = True
            context['marcoSeleccionado'] = self.request.session["seleccion"]
        else:
            context = super(MantenimientoAssessmentArchivados, self).get_context_data(**knwargs)
            context["assess"] = Assessmentguardados.objects.filter(archivado=1)
        return context


class proyectosClientes(LoginRequiredMixin, TemplateView):
    login_url = ""
    redirect_field_name = "redirect_to"
    template_name = "homepage/proyectosClientes.html"

    def get_context_data(self, **knwargs):
        context = super(proyectosClientes, self).get_context_data(**knwargs)
        context["clientes"] = Cliente.objects.all()
        context["proyectos"] = Proyecto.objects.all()
        context["usuarios"] = User.objects.all()

        return context

    def post(self, request, **knwargs):

        if 'selectorCliente' in request.POST:
            context = super(proyectosClientes, self).get_context_data(**knwargs)
            context["clientes"] = Cliente.objects.all()
            context["usuarios"] = User.objects.all()
            context["clienteSeleccionado"] = True
            context["proyectos"] = Proyecto.objects.all()
            context["cliente"] = Cliente.objects.get(codigo=request.POST.get('selectorCliente'))
            return render(request, self.template_name,
                          context=context)  # siempre retornamos el valor con la tabla completa.

        elif 'codigo' in request.POST:
            if request.POST["codigo"] != '' and request.POST["nombre"] != '':
                if not Cliente.objects.filter(codigo=request.POST["codigo"]).exists():
                    cliente_nuevo = Cliente(codigo=request.POST["codigo"],
                                            nombre=request.POST["nombre"],
                                            logo=request.FILES['logo']
                                            )
                    cliente_nuevo.save()

                else:
                    messages.error(request, 'ERROR, el Cliente ya existe de id ya existe')
            else:
                messages.error(request, 'ERROR, Necesitas introducir todos los valores')


        elif 'selectorProyecto' in request.POST:
            context = super(proyectosClientes, self).get_context_data(**knwargs)
            context["clientes"] = Cliente.objects.all()
            context["usuarios"] = User.objects.all()
            context["proyectoSeleccionado"] = True
            context["proyectos"] = Proyecto.objects.all()
            a = AsociacionUsuariosProyecto.objects.filter(
                proyecto=Proyecto.objects.get(codigo=request.POST.get('selectorProyecto')))
            p = []
            for i in a:
                p += [i.usuario.username]
            context["usuariosProyecto"] = p
            request.session['selectorProyecto'] = request.POST.get('selectorProyecto')
            context["proyecto"] = Proyecto.objects.get(codigo=request.POST.get('selectorProyecto'))
            return render(request, self.template_name,
                          context=context)  # siempre retornamos el valor con la tabla completa.
        elif 'codigoProyecto' in request.POST:
            codigo = request.POST.get('codigoProyecto')
            nombre = request.POST.get('nombreProyecto')
            descripcion = request.POST.get('descripcionProyecto')
            cliente = request.POST.get('selectorClienteProyecto')
            usuarios = request.POST.getlist('selectorUsuarios')
            if codigo != '' and nombre != '' and descripcion != '' and cliente is not None:
                if not Proyecto.objects.filter(codigo=codigo).exists():
                    cliente = Cliente.objects.get(codigo=cliente)
                    proyecto = Proyecto(codigo=codigo, nombre=nombre, descripcion=descripcion, cliente=cliente,
                                        fecha_creacion=datetime.now())
                    proyecto.save()
                    for i in usuarios:
                        user = User.objects.get(username=i)
                        asosciacion = AsociacionUsuariosProyecto(usuario=user, proyecto=proyecto)
                        asosciacion.save()
                else:
                    messages.error(request, 'ERROR, el Proyecto ya existe')
            else:
                messages.error(request, 'ERROR, Necesitas introducir todos los valores')

            context = super(proyectosClientes, self).get_context_data(**knwargs)
            context["clientes"] = Cliente.objects.all()
            context["usuarios"] = User.objects.all()
            context["proyectos"] = Proyecto.objects.all()

            return render(request, self.template_name,
                          context=context)
        elif 'selectorUsuarios2' in request.POST:
            usuarios = request.POST.getlist('selectorUsuarios2')
            proyecto = Proyecto.objects.get(codigo=request.session.get('selectorProyecto'))
            aso = AsociacionUsuariosProyecto.objects.filter(
                proyecto=Proyecto.objects.get(codigo=request.session.get('selectorProyecto')))
            aso.delete()

            for i in usuarios:
                user = User.objects.get(username=i)
                asosciacion = AsociacionUsuariosProyecto(usuario=user, proyecto=proyecto)
                asosciacion.save()

            context = super(proyectosClientes, self).get_context_data(**knwargs)
            context["clientes"] = Cliente.objects.all()
            context["usuarios"] = User.objects.all()
            context["proyectoSeleccionado"] = True
            context["proyectos"] = Proyecto.objects.all()
            a = AsociacionUsuariosProyecto.objects.filter(
                proyecto=Proyecto.objects.get(codigo=request.session.get('selectorProyecto')))
            p = []
            for i in a:
                p += [i.usuario.username]
            context["usuariosProyecto"] = p

            context["proyecto"] = Proyecto.objects.get(codigo=request.session.get('selectorProyecto'))
            return render(request, self.template_name,
                          context=context)  # siempre retornamos el valor con la tabla completa.

        context = super(proyectosClientes, self).get_context_data(**knwargs)
        context["clientes"] = Cliente.objects.all()
        context["usuarios"] = User.objects.all()
        context["proyectos"] = Proyecto.objects.all()
        return render(request, self.template_name,
                      context=context)  # siempre retornamos el valor con la tabla completa.


class entrevistasUsuarios(LoginRequiredMixin, TemplateView):
    login_url = ""
    redirect_field_name = "redirect_to"
    template_name = "homepage/EntrevistasUsuario.html"

    def get_context_data(self, **knwargs):
        context = super(entrevistasUsuarios, self).get_context_data(**knwargs)
        context["proyectos"] = AsociacionUsuariosProyecto.objects.filter(usuario=self.request.user)
        context["proyectoSelec"] = ''
        context["creadas"] = Entrevistas.objects.filter(editor=self.request.user)
        context["asistes"] = AsociacionEntrevistasUsuarios.objects.filter(usuario=self.request.user)
        return context

    def post(self, request, **knwargs):
        if 'selectorProyecto' in request.POST:
            context = super(entrevistasUsuarios, self).get_context_data(**knwargs)
            context["proyectos"] = AsociacionUsuariosProyecto.objects.filter(usuario=self.request.user)
            context["proyectoSelec"] = request.POST.get('selectorProyecto')
            context["proyectoSeleccionado"] = True
            request.session["proyectoSeleccionado"] = request.POST.get('selectorProyecto')
            context["assess"] = AsociacionProyectoAssessment.objects.filter(
                proyecto=Proyecto.objects.get(codigo=request.POST.get('selectorProyecto')), assessment__archivado=0)
            context["marcos"] = AsociacionMarcos.objects.all()
            context["creadas"] = Entrevistas.objects.filter(editor=self.request.user)
            context["asistes"] = AsociacionEntrevistasUsuarios.objects.filter(usuario=self.request.user)
            return render(request, self.template_name, context=context)

        elif 'selectorAssessment' in request.POST:
            context = super(entrevistasUsuarios, self).get_context_data(**knwargs)
            context["proyectos"] = AsociacionUsuariosProyecto.objects.filter(usuario=self.request.user)
            context["proyectoSelec"] = request.session.get('proyectoSeleccionado')
            context["proyectoSeleccionado"] = True
            request.session["proyectoSeleccionado"] = request.session.get('proyectoSeleccionado')

            context["assessSelec"] = request.POST.get('selectorAssessment')
            context["assessmentSeleccionado"] = True
            request.session["assessmentSeleccionado"] = request.POST.get('selectorAssessment')

            context["controlesAssess"] = AssessmentCreados.objects.filter(
                assessment=request.POST.get('selectorAssessment'))
            context["usuarios"] = AsociacionUsuariosProyecto.objects.filter(
                proyecto=request.session.get('proyectoSeleccionado'))
            context["assess"] = AsociacionProyectoAssessment.objects.filter(
                proyecto=Proyecto.objects.get(codigo=request.session["proyectoSeleccionado"]), assessment__archivado=0)
            context["marcos"] = AsociacionMarcos.objects.all()
            context["creadas"] = Entrevistas.objects.filter(editor=self.request.user)
            context["asistes"] = AsociacionEntrevistasUsuarios.objects.filter(usuario=self.request.user)
            return render(request, self.template_name, context=context)

        elif 'btnEditarEntrevista' in request.POST:
            entrevista = Entrevistas.objects.get(id=request.POST.get('btnEditarEntrevista'))
            context = super(entrevistasUsuarios, self).get_context_data(**knwargs)
            context["proyectos"] = AsociacionUsuariosProyecto.objects.filter(usuario=self.request.user)
            context["proyectoSelec"] = AsociacionProyectoAssessment.objects.get(
                assessment=entrevista.assesment).proyecto.codigo
            context["proyectoSeleccionado"] = True
            request.session["proyectoSeleccionado"] = AsociacionProyectoAssessment.objects.get(
                assessment=entrevista.assesment).proyecto.codigo

            context["assessSelec"] = entrevista.assesment.id_assessment
            context["assessmentSeleccionado"] = True
            request.session["assessmentSeleccionado"] = entrevista.assesment.id_assessment

            context["entrevistaEditarEditando"] = True
            context["entrevistaEditar"] = entrevista
            context["entrevistaEditarConsultores"] = AsociacionEntrevistasUsuarios.objects.filter(
                entrevista=entrevista).values_list("usuario", flat=True)
            context["entrevistaEditarFecha"] = entrevista.fecha.replace(tzinfo=None).isoformat(timespec='minutes')
            context["entrevistaEditarDuracion"] = entrevista.duracionestimada.isoformat(timespec='minutes')
            context["entrevistaControles"] = entrevista.grupocontroles.split("\n")

            context["controlesAssess"] = AssessmentCreados.objects.filter(
                assessment=entrevista.assesment.id_assessment)
            context["usuarios"] = AsociacionUsuariosProyecto.objects.filter(
                proyecto=AsociacionProyectoAssessment.objects.get(assessment=entrevista.assesment).proyecto)
            context["assess"] = AsociacionProyectoAssessment.objects.filter(
                proyecto=AsociacionProyectoAssessment.objects.get(assessment=entrevista.assesment).proyecto,
                assessment__archivado=0)
            context["marcos"] = AsociacionMarcos.objects.all()
            context["creadas"] = Entrevistas.objects.filter(editor=self.request.user)
            context["asistes"] = AsociacionEntrevistasUsuarios.objects.filter(usuario=self.request.user)
            return render(request, self.template_name, context=context)

        elif 'btnCrearEntrevista' in request.POST:

            controles = request.POST.getlist('selectorControles')
            grupoControles = ''
            for i in controles:
                grupoControles += i + '\n'
            if request.POST.get('Titulo') != '' and request.POST.get('Fecha') != '' and request.POST.get(
                    'Area') != '' and request.POST.get('Duracion') != '' and request.POST.get(
                'selectorEditor') != '' and grupoControles != '':
                entrevista = Entrevistas(
                    titulo=request.POST.get('Titulo'),
                    fecha=datetime.fromisoformat(request.POST.get('Fecha')),
                    grupocontroles=grupoControles,
                    area=request.POST.get('Area'),
                    creador=User.objects.get(username=request.user),
                    duracionestimada=request.POST.get('Duracion'),
                    assesment=Assessmentguardados.objects.get(
                        id_assessment=request.session.get('assessmentSeleccionado')),
                    editor=User.objects.get(username=request.POST.get('selectorEditor')),
                    asistentes=request.POST.get('Asistentes')
                )
                entrevista.save()
                users = request.POST.getlist('selectorUsuarios')
                for i in users:
                    usuario = User.objects.get(username=i)
                    asociar = AsociacionEntrevistasUsuarios(entrevista=entrevista, usuario=usuario)
                    asociar.save()
            else:
                messages.error(request, 'ERROR, Necesitas introducir todos los valores')
            context = super(entrevistasUsuarios, self).get_context_data(**knwargs)
            context["proyectos"] = AsociacionUsuariosProyecto.objects.filter(usuario=self.request.user)

            context["creadas"] = Entrevistas.objects.filter(editor=self.request.user)
            context["asistes"] = AsociacionEntrevistasUsuarios.objects.filter(usuario=self.request.user)
            return render(request, self.template_name, context=context)

        elif 'btnEditarEntrevistaEnviar' in request.POST:

            controles = request.POST.getlist('selectorControles')
            grupoControles = ''
            for i in controles:
                grupoControles += i + '\n'
            if request.POST.get('Titulo') != '' and request.POST.get('Fecha') != '' and request.POST.get(
                    'Area') != '' and request.POST.get('Duracion') != '' and request.POST.get(
                'selectorEditor') != '' and grupoControles != '':
                entrevista = Entrevistas.objects.get(id=request.POST.get('btnEditarEntrevistaEnviar'))

                entrevista.titulo = request.POST.get('Titulo')
                entrevista.fecha = datetime.fromisoformat(request.POST.get('Fecha'))
                entrevista.grupocontroles = grupoControles
                entrevista.area = request.POST.get('Area')
                entrevista.creador = User.objects.get(username=request.user)
                entrevista.duracionestimada = request.POST.get('Duracion')
                entrevista.assesment = Assessmentguardados.objects.get(
                    id_assessment=request.session.get('assessmentSeleccionado'))
                entrevista.editor = User.objects.get(username=request.POST.get('selectorEditor'))
                entrevista.asistentes = request.POST.get('Asistentes')

                entrevista.save()
                users = request.POST.getlist('selectorUsuarios')
                for i in users:
                    usuario = User.objects.get(username=i)
                    asociar = AsociacionEntrevistasUsuarios(entrevista=entrevista, usuario=usuario)
                    asociar.save()
            else:
                messages.error(request, 'ERROR, Necesitas introducir todos los valores')
            context = super(entrevistasUsuarios, self).get_context_data(**knwargs)
            context["proyectos"] = AsociacionUsuariosProyecto.objects.filter(usuario=self.request.user)

            context["creadas"] = Entrevistas.objects.filter(editor=self.request.user)
            context["asistes"] = AsociacionEntrevistasUsuarios.objects.filter(usuario=self.request.user)
            return render(request, self.template_name, context=context)

        elif 'btnEditarAssesment' in request.POST:
            request.session["EntrevistaEditar"] = request.POST.get('btnEditarAssesment')
            return redirect("encuestaEntrevista")
        elif 'btnEliminarEntrevista' in request.POST:
            entrevista = Entrevistas.objects.get(id=request.POST.get('btnEliminarEntrevista'))
            entrevista.delete()
            context = super(entrevistasUsuarios, self).get_context_data(**knwargs)
            context["proyectos"] = AsociacionUsuariosProyecto.objects.filter(usuario=self.request.user)
            context["proyectoSelec"] = ''
            context["creadas"] = Entrevistas.objects.filter(editor=self.request.user)
            context["asistes"] = AsociacionEntrevistasUsuarios.objects.filter(usuario=self.request.user)
            return render(request, self.template_name, context=context)


class planProyecto(LoginRequiredMixin, TemplateView):
    login_url = ""
    redirect_field_name = "redirect_to"
    template_name = "homepage/planProyecto.html"

    def contexto(self, context):
        assSelect = self.request.session.get('assessmentGuardado')
        ass = Assessmentguardados.objects.get(id_assessment=assSelect)
        if self.request.session["ProyectoSeleccionado"] != None:
            if ProyectosMejora.objects.filter(id=self.request.session["ProyectoSeleccionado"]).exists():
                plan = ProyectosMejora.objects.get(id=self.request.session["ProyectoSeleccionado"])

                g = []
                c = []
                for i in plan.descripcion.split(" "):
                    p = AsociacionEvidenciasGenericas.objects.filter(assessment__assessment=ass,
                                                                     iniciativa__descripcion__icontains=i)
                    for s in p:
                        if s not in g and not AsociacionProyectoMejoraIniciativa.objects.filter(proyecto=plan,
                                                                                                iniciativa=s.iniciativa).exists():
                            g += [s]
                    p = AsociacionEvidenciasCreadas.objects.filter(id_assessment__assessment=ass,
                                                                   iniciativa__descripcion__icontains=i)
                    for s in p:
                        if s not in c and not AsociacionProyectoMejoraIniciativa.objects.filter(proyecto=plan,
                                                                                                iniciativa=s.iniciativa).exists():
                            c += [s]
                context["recomendacion"] = g + c
                context["asociadas"] = AsociacionProyectoMejoraIniciativa.objects.filter(proyecto=plan)
                asoci = []
                for i in AsociacionProyectoMejoraIniciativa.objects.filter(proyecto=plan):
                    asoci += [i.iniciativa.id]
                context["idAsoci"] = asoci
            context["ProyectoSeleccionado"] = ProyectosMejora.objects.get(
                id=self.request.session["ProyectoSeleccionado"])
        else:
            context["ProyectoSeleccionado"] = self.request.session["ProyectoSeleccionado"]
        if ass.plan_proyecto_mejora != None:
            context["plan"] = ass.plan_proyecto_mejora
        context["proyectos"] = AsociacionPlanProyectosProyectos.objects.filter(plan_proyecto=ass.plan_proyecto_mejora)
        if AsociacionPlanProyectosProyectos.objects.filter(plan_proyecto=ass.plan_proyecto_mejora).exists():
            context["primero"] = \
                AsociacionPlanProyectosProyectos.objects.filter(plan_proyecto=ass.plan_proyecto_mejora)[
                    0].proyecto_mejora.id

        context["iniciativas"] = AsociacionEvidenciasGenericas.objects.filter(assessment__assessment=ass)
        context["iniciativasc"] = AsociacionEvidenciasCreadas.objects.filter(id_assessment__assessment=ass)
        context["assess"] = Assessmentguardados.objects.get(id_assessment=assSelect)
        return context

    def get_context_data(self, **knwargs):
        self.request.session["ProyectoSeleccionado"] = None
        context = super(planProyecto, self).get_context_data(**knwargs)
        context = self.contexto(context)
        return context

    def post(self, request, **knwargs):
        assSelect = request.session.get('assessmentGuardado')
        if 'crearProyecto' in request.POST:
            if request.POST.get('NombrePlan') != '' and request.POST.get('descripcionPlan') != '' and request.POST.get(
                    'riesgosPlan') != '' and request.POST.get('tipoPlan') != '' and request.POST.get(
                'duracionPlan') != '' and request.POST.get('costePlan') != '' and request.POST.get(
                'beneficioPlan') != '':
                ass = Assessmentguardados.objects.get(id_assessment=assSelect)
                if ass.plan_proyecto_mejora != None:
                    proyecto = ProyectosMejora(nombre=request.POST.get('NombrePlan'),
                                               descripcion=request.POST.get('descripcionPlan'),
                                               riesgos=request.POST.get('riesgosPlan'),
                                               tipo=request.POST.get('tipoPlan'),
                                               duracion=request.POST.get('duracionPlan'),
                                               coste=request.POST.get('costePlan'),
                                               beneficio=request.POST.get('beneficioPlan'),
                                               )
                    proyecto.save()
                    plan = ass.plan_proyecto_mejora
                    asociacion = AsociacionPlanProyectosProyectos(proyecto_mejora=proyecto, plan_proyecto=plan)
                    asociacion.save()

                    context = super(planProyecto, self).get_context_data(**knwargs)
                    context = self.contexto(context)

                    return render(request, self.template_name, context=context)
                else:
                    messages.error(request, 'ERROR, Necesitas crear un plan de proyecto antes de crear un proyecto')
            else:
                messages.error(request, 'ERROR, Necesitas introducir todos los valores')

            context = super(planProyecto, self).get_context_data(**knwargs)
            context = self.contexto(context)
            return render(request, self.template_name, context=context)
        elif 'editarProyecto' in request.POST:
            if request.POST.get('NombrePlan') != '' and request.POST.get('descripcionPlan') != '' and request.POST.get(
                    'riesgosPlan') != '' and request.POST.get('tipoPlan') != '' and request.POST.get(
                'duracionPlan') != '' and request.POST.get('capex') != '' and request.POST.get(
                'beneficioPlan') != '' and request.POST.get('opex') != '':
                ass = Assessmentguardados.objects.get(id_assessment=assSelect)
                proyecto = ProyectosMejora.objects.get(id=request.session["ProyectoEditar"])
                proyecto.nombre = request.POST.get('NombrePlan')
                proyecto.descripcion = request.POST.get('descripcionPlan')
                proyecto.riesgos = request.POST.get('riesgosPlan')
                proyecto.tipo = request.POST.get('tipoPlan')
                proyecto.duracion = request.POST.get('duracionPlan')
                proyecto.capex = request.POST.get('capex')
                proyecto.opex = request.POST.get('opex')
                proyecto.beneficio = request.POST.get('beneficioPlan')
                proyecto.save()

                context = super(planProyecto, self).get_context_data(**knwargs)
                context = self.contexto(context)

                return render(request, self.template_name, context=context)

            else:
                messages.error(request, 'ERROR, Necesitas introducir todos los valores')

            context = super(planProyecto, self).get_context_data(**knwargs)
            context = self.contexto(context)
            return render(request, self.template_name, context=context)
        elif 'NombrePlanProyecto' in request.POST:
            ass = Assessmentguardados.objects.get(id_assessment=assSelect)
            if ass.plan_proyecto_mejora == None:
                plan = PlanProyectoMejora(nombre=request.POST.get("NombrePlanProyecto"),
                                          descripcion=request.POST.get("descripcionPlanProyecto"))
                plan.save()
                ass = Assessmentguardados.objects.get(id_assessment=assSelect)
                ass.plan_proyecto_mejora = plan
                ass.save()
            else:
                plan = ass.plan_proyecto_mejora
                plan.nombre = request.POST.get("NombrePlanProyecto")
                plan.descripcion = request.POST.get("descripcionPlanProyecto")
                plan.save()
            context = super(planProyecto, self).get_context_data(**knwargs)
            context = self.contexto(context)
            return render(request, self.template_name, context=context)
        elif 'btnSeleccionProyecto' in request.POST:
            request.session["ProyectoSeleccionado"] = request.POST.get('btnSeleccionProyecto')

            context = super(planProyecto, self).get_context_data(**knwargs)
            context = self.contexto(context)
            return render(request, self.template_name, context=context)

        elif 'btnEditarProyecto' in request.POST:
            request.session["ProyectoEditar"] = request.POST.get('btnEditarProyecto')

            context = super(planProyecto, self).get_context_data(**knwargs)
            context["ProyectoEditar"] = ProyectosMejora.objects.get(id=request.session["ProyectoEditar"])
            context["duracion"] = str(ProyectosMejora.objects.get(id=request.session["ProyectoEditar"]).duracion)
            context["capex"] = str(ProyectosMejora.objects.get(id=request.session["ProyectoEditar"]).capex)
            context["opex"] = str(ProyectosMejora.objects.get(id=request.session["ProyectoEditar"]).opex)
            context["beneficio"] = str(ProyectosMejora.objects.get(id=request.session["ProyectoEditar"]).beneficio)
            context["editandoProyecto"] = True
            context = self.contexto(context)
            return render(request, self.template_name, context=context)

        elif 'selectorIniciativasg' in request.POST:
            plan = ProyectosMejora.objects.get(id=request.session["ProyectoSeleccionado"])
            iniciaticasSelect = request.POST.getlist('selectorIniciativasg')
            for i in iniciaticasSelect:
                if not AsociacionProyectoMejoraIniciativa.objects.filter(proyecto=plan,
                                                                         iniciativa=Iniciativas.objects.get(
                                                                             nombre=i)).exists():
                    asociacion = AsociacionProyectoMejoraIniciativa(proyecto=plan,
                                                                    iniciativa=Iniciativas.objects.get(nombre=i))
                    asociacion.save()
            context = super(planProyecto, self).get_context_data(**knwargs)
            context = self.contexto(context)
            return render(request, self.template_name, context=context)

        elif 'selectorIniciativasc' in request.POST:
            plan = ProyectosMejora.objects.get(id=request.session["ProyectoSeleccionado"])
            iniciaticasSelect = request.POST.getlist('selectorIniciativasc')

            for i in iniciaticasSelect:
                if not AsociacionProyectoMejoraIniciativa.objects.filter(proyecto=plan,
                                                                         iniciativa=Iniciativas.objects.get(
                                                                             nombre=i)).exists():
                    asociacion = AsociacionProyectoMejoraIniciativa(proyecto=plan,
                                                                    iniciativa=Iniciativas.objects.get(nombre=i))
                    asociacion.save()
            context = super(planProyecto, self).get_context_data(**knwargs)
            context = self.contexto(context)
            return render(request, self.template_name, context=context)

        elif 'btnAnadirRecomendacion' in request.POST:
            plan = ProyectosMejora.objects.get(id=request.session["ProyectoSeleccionado"])
            iniciaticasSelect = request.POST.get('btnAnadirRecomendacion')
            if not AsociacionProyectoMejoraIniciativa.objects.filter(proyecto=plan, iniciativa=Iniciativas.objects.get(
                    nombre=iniciaticasSelect)).exists():
                asociacion = AsociacionProyectoMejoraIniciativa(proyecto=plan,
                                                                iniciativa=Iniciativas.objects.get(
                                                                    nombre=iniciaticasSelect))
                asociacion.save()
            context = super(planProyecto, self).get_context_data(**knwargs)
            context = self.contexto(context)
            return render(request, self.template_name, context=context)

        elif 'btnEliminarAsociacion' in request.POST:
            plan = ProyectosMejora.objects.get(id=request.session["ProyectoSeleccionado"])
            iniciaticasSelect = request.POST.get('btnEliminarAsociacion')
            if AsociacionProyectoMejoraIniciativa.objects.filter(proyecto=plan, iniciativa=Iniciativas.objects.get(
                    nombre=iniciaticasSelect)).exists():
                asociacion = AsociacionProyectoMejoraIniciativa.objects.get(proyecto=plan,
                                                                            iniciativa=Iniciativas.objects.get(
                                                                                nombre=iniciaticasSelect))
                asociacion.delete()
            context = super(planProyecto, self).get_context_data(**knwargs)
            context = self.contexto(context)
            return render(request, self.template_name, context=context)


# Clase para la pagina de Entrevista
class encuestaEntrevista(LoginRequiredMixin, TemplateView):
    login_url = ""
    redirect_field_name = "redirect_to"
    template_name = "homepage/EncuestaEntrevista.html"

    def contextTotal(self, request, select, context):

        entre = self.request.session.get('EntrevistaEditar')
        entrevista = Entrevistas.objects.get(id=entre)
        assSelect = entrevista.assesment.id_assessment
        assGuardado = entrevista.assesment
        e = entrevista.grupocontroles.split('\n')
        e = e[:len(e) - 1]

        context["primero"] = False
        context["ultimo"] = False
        if e.index(select) == len(e) - 1:
            context["ultimo"] = True
        elif e.index(select) == 0:
            context["primero"] = True

        context["controlEntrevista"] = e[:len(e)]
        context["NombreAss"] = assSelect
        context["assess"] = AssessmentCreados.objects.filter(assessment=assGuardado)
        control = AssessmentCreados.objects.get(assessment=assGuardado, control_id=select)
        context["control"] = control
        self.request.session["controlSelect"] = select
        context["evidencias"] = AsociacionEvidenciasGenericas.objects.filter(assessment=control)
        context["evidencias2"] = AsociacionEvidenciasCreadas.objects.filter(id_assessment=control)
        if assGuardado.idioma == 'en':
            context["evidenciasGenerricas"] = Evidencerequestcatalog.objects.all()
            lista = []
            for i in AsociacionEvidenciasGenericas.objects.filter(assessment=control):
                lista += [i.evidencia]
            context["listaEvidencias"] = lista
        else:
            context["evidenciasGenerricas"] = EvidencerequestcatalogEs.objects.all()
            lista = []
            for i in AsociacionEvidenciasGenericas.objects.filter(assessment=control):
                lista += [i.evidencia_id_es]
            context["listaEvidencias"] = lista
        return request, context

    # funcion que envia el contexto de la pagina.
    def get_context_data(self, **knwargs):
        entre = self.request.session.get('EntrevistaEditar')
        entrevista = Entrevistas.objects.get(id=entre)
        context = super(encuestaEntrevista, self).get_context_data(**knwargs)
        assSelect = entrevista.assesment.id_assessment
        assGuardado = entrevista.assesment
        e = entrevista.grupocontroles.split('\n')
        e = e[:len(e) - 1]
        context["controlEntrevista"] = e
        context["ultimo"] = False
        context["primero"] = True
        context["NombreAss"] = assSelect
        context["assess"] = AssessmentCreados.objects.filter(assessment=assGuardado)
        control = AssessmentCreados.objects.get(assessment=assGuardado, control_id=e[0])
        if assGuardado.idioma == 'en':
            context["evidenciasGenerricas"] = Evidencerequestcatalog.objects.all()
            lista = []
            for i in AsociacionEvidenciasGenericas.objects.filter(assessment=control):
                lista += [i.evidencia]
            context["listaEvidencias"] = lista
        else:
            context["evidenciasGenerricas"] = EvidencerequestcatalogEs.objects.all()
            lista = []
            for i in AsociacionEvidenciasGenericas.objects.filter(assessment=control):
                lista += [i.evidencia_id_es]
            context["listaEvidencias"] = lista

        context["control"] = control
        self.request.session["controlSelect"] = e[0]
        context["evidencias"] = AsociacionEvidenciasGenericas.objects.filter(assessment=control)
        context["evidencias2"] = AsociacionEvidenciasCreadas.objects.filter(id_assessment=control)
        return context

    # funcion post que recoge los summit del formulario de la pagina.
    def post(self, request, **knwargs):
        entre = request.session.get('EntrevistaEditar')
        entrevista = Entrevistas.objects.get(id=entre)
        assSelect = entrevista.assesment.id_assessment
        select = request.POST.get('selector')  # valor de el selector de control
        boton2 = request.POST.get('boton2')  # valor del boton 2
        boton3 = request.POST.get('boton3')  # valor del boton 3
        boton4 = request.POST.get('boton4')  # valor del boton 4
        boton5 = request.POST.get('boton5')  # valor del boton 5
        boton6 = request.POST.get('boton6')  # valor del boton 6
        boton7 = request.POST.get('boton7')  # valor del boton 7
        btnEliminarEvidencia = request.POST.get('btnEliminarEvidencia')  # valor del btnEliminarEvidencia

        if 'selector' in request.POST:  # se recoge la pulsacion del select
            if select == 'noSel':
                entrevista = Entrevistas.objects.get(id=entre)
                context = super(encuestaEntrevista, self).get_context_data(**knwargs)
                assSelect = entrevista.assesment.id_assessment
                assGuardado = entrevista.assesment
                e = entrevista.grupocontroles.split('\n')
                context["controlEntrevista"] = e[:len(e) - 1]
                context["NombreAss"] = assSelect
                context["assess"] = AssessmentCreados.objects.filter(assessment=assGuardado)
                self.request.session["controlSelect"] = 'noSel'
                return render(request, self.template_name, context=context)
            else:
                context = super(encuestaEntrevista, self).get_context_data(**knwargs)
                request, context = self.contextTotal(request, select, context)
                return render(request, self.template_name, context=context)

        elif boton2 == 'btn2':  # recogemos la pulsacion del boton de guardar valoracion
            controlSeleccionado = request.session.get('controlSelect')
            if request.session["controlSelect"] != 'noSel':
                assGuardado = Assessmentguardados.objects.get(id_assessment=assSelect)
                control = AssessmentCreados.objects.get(assessment=assGuardado,
                                                        control_id=request.session["controlSelect"])
                control.respuesta = str(request.POST.get('respuesta'))
                control.valoracion = request.POST.get('valmad')
                control.valoracionobjetivo = request.POST.get('valmadob')
                control.save()
                e = entrevista.grupocontroles.split('\n')
                e = e[:len(e) - 1]
                controlSeleccionado = e[e.index(request.session.get('controlSelect')) + 1]

            else:
                messages.error(request,
                               'CONTROL INCORRECTO: Necesita seleccionar un control para realizar esta acción')  # Se crea mensage de error
            context = super(encuestaEntrevista, self).get_context_data(**knwargs)
            request, context = self.contextTotal(request, controlSeleccionado, context)
            return render(request, self.template_name, context=context)

        elif boton6 == 'btn6':  # recogemos la pulsacion del boton de guardar valoracion
            controlSeleccionado = request.session.get('controlSelect')
            if request.session["controlSelect"] != 'noSel':
                assGuardado = Assessmentguardados.objects.get(id_assessment=assSelect)
                control = AssessmentCreados.objects.get(assessment=assGuardado,
                                                        control_id=request.session["controlSelect"])
                control.respuesta = str(request.POST.get('respuesta'))
                control.valoracion = request.POST.get('valmad')
                control.valoracionobjetivo = request.POST.get('valmadob')
                control.save()
                e = entrevista.grupocontroles.split('\n')
                e = e[:len(e) - 1]
                controlSeleccionado = e[e.index(request.session.get('controlSelect')) - 1]

            else:
                messages.error(request,
                               'CONTROL INCORRECTO: Necesita seleccionar un control para realizar esta acción')  # Se crea mensage de error
            context = super(encuestaEntrevista, self).get_context_data(**knwargs)
            request, context = self.contextTotal(request, controlSeleccionado, context)
            return render(request, self.template_name, context=context)

        elif boton7 == 'btn7':  # recogemos la pulsacion del boton de guardar valoracion
            controlSeleccionado = request.session.get('controlSelect')
            if request.session["controlSelect"] != 'noSel':
                assGuardado = Assessmentguardados.objects.get(id_assessment=assSelect)
                control = AssessmentCreados.objects.get(assessment=assGuardado,
                                                        control_id=request.session["controlSelect"])
                control.respuesta = str(request.POST.get('respuesta'))
                control.valoracion = request.POST.get('valmad')
                control.valoracionobjetivo = request.POST.get('valmadob')
                control.save()
                controlSeleccionado = request.session.get('controlSelect')
            else:
                messages.error(request,
                               'CONTROL INCORRECTO: Necesita seleccionar un control para realizar esta acción')  # Se crea mensage de error
            context = super(encuestaEntrevista, self).get_context_data(**knwargs)
            request, context = self.contextTotal(request, controlSeleccionado, context)
            return render(request, self.template_name, context=context)

        elif boton3 == 'btn3':  # recogemos la pulsacion del boton de guardar valoracion
            controlSeleccionado = request.session.get('controlSelect')
            if request.session["controlSelect"] != 'noSel':
                assGuardado = Assessmentguardados.objects.get(id_assessment=assSelect)
                control = AssessmentCreados.objects.get(assessment=assGuardado,
                                                        control_id=request.session["controlSelect"])
                control.respuesta = str(request.POST.get('respuesta'))
                control.valoracion = request.POST.get('valmad')
                control.valoracionobjetivo = request.POST.get('valmadob')
                control.save()


            else:
                messages.error(request,
                               'CONTROL INCORRECTO: Necesita seleccionar un control para realizar esta acción')  # Se crea mensage de error

            return redirect('entrevistasUsuarios')
        elif boton4 == 'btn4':  # if encargado de rellenar las evidencias

            idEvidencia = request.POST.get('idEvidencia') + '-C'  # valor del idEvidencia
            descripcionEvidencia = request.POST.get('DescripcionEvidencia')  # valor del DescripcionEvidencia
            linkEvidencia = request.POST.get('linkEvidencia')  # valor del linkEvidencia
            controlId = request.session["controlSelect"]

            if controlId != 'noSel':
                if idEvidencia != '' and descripcionEvidencia != '':  # recogemos la pulsacion de guardar la evidencia
                    if not Evidencias.objects.filter(evidencia_id=idEvidencia).exists():

                        ev = Evidencias(evidencia_id=idEvidencia, comentario=descripcionEvidencia, links=linkEvidencia,
                                        control_id=controlId,
                                        assessment=Assessmentguardados.objects.get(id_assessment=assSelect))
                        ev.save()
                        assGuardado = Assessmentguardados.objects.get(id_assessment=assSelect)
                        control = AssessmentCreados.objects.get(assessment=assGuardado,
                                                                control_id=request.session["controlSelect"])
                        control.respuesta = str(request.POST.get('respuesta'))
                        control.valoracion = request.POST.get('valmad')
                        control.valoracionobjetivo = request.POST.get('valmadob')

                        evidencia = AsociacionEvidenciasCreadas(id_evidencia=ev, id_assessment=control)
                        evidencia.save()
                        context = super(encuestaEntrevista, self).get_context_data(**knwargs)
                        request, context = self.contextTotal(request, controlId, context)
                        return render(request, self.template_name, context=context)
                    else:
                        context = super(encuestaEntrevista, self).get_context_data(**knwargs)
                        request, context = self.contextTotal(request, controlId, context)
                        messages.error(request,
                                       'EVIDENCIA INCORRECTA: La evidencia introducida ya existe')  # Se crea mensage de error
                        return render(request, self.template_name, context=context)
                else:
                    context = super(encuestaEntrevista, self).get_context_data(**knwargs)
                    request, context = self.contextTotal(request, controlId, context)
                    messages.error(request, 'EVIDENCIA INCORRECTA Necesita introducir un id y una descripcion para la '
                                            'evidencia')  # Se crea mensage de error
                    return render(request, self.template_name, context=context)
            else:
                messages.error(request,
                               'CONTROL INCORRECTO: Necesita seleccionar un control para realizar esta acción')  # Se crea mensage de error
                context = super(encuestaEntrevista, self).get_context_data(**knwargs)
                request, context = self.contextTotal(request, controlId, context)
                return render(request, self.template_name, context=context)

        elif boton5 == 'btn5':  # if encargado de rellenar las evidencias
            if request.session["controlSelect"] != 'noSel':
                selectorEvidencia = request.POST.get('selectorEvidencia')  # valor de el selector de control
                if selectorEvidencia != 'noSel':
                    assGuardado = Assessmentguardados.objects.get(id_assessment=assSelect)
                    control = AssessmentCreados.objects.get(assessment=assGuardado,
                                                            control_id=request.session["controlSelect"])
                    if assGuardado.idioma == 'en':
                        evidencia = Evidencerequestcatalog.objects.get(evidence_request_references=selectorEvidencia)
                        eviGenerica = AsociacionEvidenciasGenericas(evidencia=evidencia, assessment=control)
                        eviGenerica.save()
                    else:
                        evidencia = EvidencerequestcatalogEs.objects.get(evidence_request_references=selectorEvidencia)
                        eviGenerica = AsociacionEvidenciasGenericas(evidencia_id_es=evidencia, assessment=control)
                        eviGenerica.save()
                    assGuardado = Assessmentguardados.objects.get(id_assessment=assSelect)
                    control = AssessmentCreados.objects.get(assessment=assGuardado,
                                                            control_id=request.session["controlSelect"])
                    control.respuesta = str(request.POST.get('respuesta'))
                    control.valoracion = request.POST.get('valmad')
                    control.valoracionobjetivo = request.POST.get('valmadob')
                    control.save()
                    context = super(encuestaEntrevista, self).get_context_data(**knwargs)
                    request, context = self.contextTotal(request, request.session["controlSelect"], context)
                    return render(request, self.template_name, context=context)
                else:
                    context = super(encuestaEntrevista, self).get_context_data(**knwargs)
                    request, context = self.contextTotal(request, request.session["controlSelect"], context)
                    messages.error(request,
                                   'EVIDENCIA INCORRECTA: Necesita seleccionar una evidencia')  # Se crea mensage de error
                    return render(request, self.template_name, context=context)
            else:
                messages.error(request,
                               'CONTROL INCORRECTO: Necesita seleccionar un control para realizar esta acción')  # Se crea mensage de error
                context = super(encuestaEntrevista, self).get_context_data(**knwargs)
                request, context = self.contextTotal(request, request.session["controlSelect"], context)
                return render(request, self.template_name, context=context)

        elif btnEliminarEvidencia != '':  # if encargado de Eliminar las evidencias
            evidencia = AsociacionEvidenciasGenericas.objects.get(id=btnEliminarEvidencia)
            evidencia.delete()
            context = super(encuestaEntrevista, self).get_context_data(**knwargs)
            request, context = self.contextTotal(request, request.session["controlSelect"], context)
            return render(request, self.template_name, context=context)
